#!/usr/bin/env python3
"""
Weekly Market Intelligence Report Generator
============================================
Pulls 7 days of data from Supabase + Google News, sends to Claude Opus,
produces both long-form (~1500 words) and short-form (~600 words) reports
in Markdown and DOCX. News headlines are first run through a Claude Haiku
event-extraction pre-pass that digests the raw feed into a categorised
list of market-moving events (the report's authoritative source for "why").

Schedule: Saturday 03:00 UTC = 08:30 IST (after Friday close,
after Sat 7 AM Screener fetch).

Outputs:
  reports/YYYY-MM-DD_market-intel_long.md
  reports/YYYY-MM-DD_market-intel_long.docx
  reports/YYYY-MM-DD_market-intel_short.md
  reports/YYYY-MM-DD_market-intel_short.docx
  reports/_context_YYYY-MM-DD.md   (raw LLM context, for debugging)
"""

import os
import re
import json
import html
import urllib.parse
from collections import defaultdict
from datetime import datetime, timedelta, date

import feedparser
import requests
from supabase import create_client
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt


# =============================================================================
# CONFIG
# =============================================================================

SUPABASE_URL = os.environ.get('SUPABASE_URL', 'https://hcgyncghmcvylnrmcivj.supabase.co')
SUPABASE_KEY = os.environ['SUPABASE_KEY']
ANTHROPIC_API_KEY = os.environ['ANTHROPIC_API_KEY']

MODEL = 'claude-opus-4-7'
HAIKU_MODEL = 'claude-haiku-4-5-20251001'   # news event-extraction pre-pass
OUT_DIR = 'reports'

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
claude = Anthropic(api_key=ANTHROPIC_API_KEY)


# =============================================================================
# SUPABASE FETCHERS
# =============================================================================

def paginated_fetch(table, select, filters_fn=None, batch=1000, max_rows=10000):
    out, off = [], 0
    while off < max_rows:
        q = supabase.table(table).select(select)
        if filters_fn:
            q = filters_fn(q)
        try:
            resp = q.range(off, off + batch - 1).execute()
        except Exception as e:
            print(f"    paginated_fetch error on {table}: {e}")
            return out
        if not resp.data:
            break
        out.extend(resp.data)
        if len(resp.data) < batch:
            break
        off += batch
    return out


def fetch_macro_snapshot():
    """Latest market_ratios row + week-ago for deltas.
    Note: Nifty/Sensex are NOT pulled from here — market_ratios_scanner runs
    at 7 AM IST before market opens, so its Indian-index values are stale.
    Use fetch_daily_index_closes() for those instead."""
    today = date.today()
    floor = (today - timedelta(days=12)).isoformat()
    rows = paginated_fetch(
        'market_ratios', '*',
        lambda q: q.gte('snapshot_date', floor).order('snapshot_date', desc=True),
    )
    if not rows:
        return None
    latest = rows[0]
    target = today - timedelta(days=7)
    week_ago = min(rows, key=lambda r: abs((date.fromisoformat(r['snapshot_date']) - target).days))

    fields = [
        ('india_vix', 'India VIX', None),
        ('crude_oil', 'Brent Crude (USD)', '%'),
        ('dxy', 'US Dollar Index (DXY)', None),
        ('gold_price', 'Gold (USD/oz)', '%'),
        ('silver_price', 'Silver (USD/oz)', '%'),
        ('gold_silver_ratio', 'Gold/Silver Ratio', None),
        ('nifty_pe', 'Nifty P/E', None),
        ('market_cap_gdp_ratio', 'Market Cap / GDP', None),
        ('sp500', 'S&P 500', '%'),
        ('nasdaq', 'Nasdaq', '%'),
    ]
    items = []
    for key, label, fmt in fields:
        cur = latest.get(key)
        prev = week_ago.get(key)
        if cur is None:
            continue
        item = {'label': label, 'current': cur}
        if prev is not None and prev != 0:
            if fmt == '%':
                item['change_pct'] = (cur - prev) / prev * 100
            else:
                item['change_abs'] = cur - prev
        items.append(item)
    return {
        'latest_date': latest['snapshot_date'],
        'week_ago_date': week_ago['snapshot_date'],
        'items': items,
    }


def compute_ema(series, span):
    """Exponential moving average — current value after processing the full series.
    Matches pandas .ewm(span=N).mean().iloc[-1] when adjust=False.
    Returns None for series shorter than span."""
    if not series or len(series) < span:
        return None
    alpha = 2.0 / (span + 1)
    ema_val = series[0]
    for v in series[1:]:
        ema_val = (v - ema_val) * alpha + ema_val
    return ema_val


def fetch_nifty_emas():
    """Pulls last ~280 days of nifty_50 from market_ratios and computes EMA20/50/200.
    market_ratios is 1 day stale on the latest value but for 200-day EMA computation
    that's irrelevant. Current Nifty close should be cross-referenced against
    detailed_snapshots (fetch_daily_index_closes) for the freshest value."""
    today = date.today()
    floor = (today - timedelta(days=400)).isoformat()
    rows = paginated_fetch(
        'market_ratios', 'snapshot_date, nifty_50',
        lambda q: q.gte('snapshot_date', floor).not_.is_('nifty_50', 'null')
                   .order('snapshot_date', desc=False),
    )
    if not rows or len(rows) < 60:
        return None
    closes = [r['nifty_50'] for r in rows if r['nifty_50'] is not None]
    if len(closes) < 60:
        return None
    return {
        'last_close':       round(closes[-1], 2),
        'last_close_date':  rows[-1]['snapshot_date'],
        'ema_20':           round(compute_ema(closes, 20),  2) if len(closes) >= 20  else None,
        'ema_50':           round(compute_ema(closes, 50),  2) if len(closes) >= 50  else None,
        'ema_200':          round(compute_ema(closes, 200), 2) if len(closes) >= 200 else None,
        'days_used':        len(closes),
    }


def fetch_daily_index_closes():
    """Last 5–7 trading days of NIFTY / SENSEX / NIFTYBEES from detailed_snapshots.
    This is the authoritative source for accurate daily index closes
    (written 3:40 PM IST after market close, unlike market_ratios which is 7 AM).
    Returns chronologically ordered list with day-of-week labels."""
    try:
        resp = supabase.table('detailed_snapshots') \
            .select('snapshot_date, indices') \
            .eq('portfolio', 'INDIAN') \
            .order('snapshot_date', desc=True).limit(7).execute()
    except Exception as e:
        print(f"    fetch_daily_index_closes error: {e}")
        return []
    out = []
    for r in (resp.data or []):
        d = r['snapshot_date']
        idx = r.get('indices') or {}
        if isinstance(idx, str):
            try:
                idx = json.loads(idx)
            except Exception:
                idx = {}
        out.append({
            'date': d,
            'day':  datetime.strptime(d, '%Y-%m-%d').strftime('%A'),
            'nifty':     idx.get('NIFTY'),
            'sensex':    idx.get('SENSEX'),
            'niftybees': idx.get('NIFTYBEES'),
        })
    out = list(reversed(out))  # chronological — Mon → Fri

    # Compute weekly deltas from first to last entry
    summary = {'days': out}
    if len(out) >= 2:
        first, last = out[0], out[-1]
        for k in ('nifty', 'sensex', 'niftybees'):
            f, l = first.get(k), last.get(k)
            if f and l and f != 0:
                summary[f'{k}_pct'] = (l - f) / f * 100
                summary[f'{k}_abs'] = l - f
                summary[f'{k}_close'] = l
                summary[f'{k}_first'] = f
        summary['date_first'] = first['date']
        summary['date_last']  = last['date']
    return summary


def fetch_sector_summary():
    """Top/bottom 5 sectors by 5-day return from view_sector_performance."""
    rows = paginated_fetch('view_sector_performance', '*')
    if not rows:
        return None
    rows = [r for r in rows if r.get('med_5d') is not None]
    rows.sort(key=lambda r: r['med_5d'], reverse=True)
    return {
        'top_5': rows[:5],
        'bottom_5': rows[-5:][::-1],
        'all_count': len(rows),
    }


def quadrant_of(rs, mom):
    if rs >= 50 and mom >= 0:
        return 'Leading'
    if rs >= 50 and mom < 0:
        return 'Weakening'
    if rs < 50 and mom < 0:
        return 'Lagging'
    return 'Improving'


def fetch_quadrant_flips():
    """Sectors that flipped quadrants between today and 7 days ago.
    Computes momentum as 10-day slope of median_rs at each anchor date."""
    rows = paginated_fetch(
        'view_sector_rs_history', 'sector, snapshot_date, median_rs',
        lambda q: q.order('snapshot_date', desc=False),
    )
    if not rows:
        return []

    by_sector = defaultdict(list)
    for r in rows:
        by_sector[r['sector']].append(r)

    flips = []
    for sector, history in by_sector.items():
        history.sort(key=lambda r: r['snapshot_date'])
        if len(history) < 18:
            continue
        # Today's quadrant
        today_rs = history[-1]['median_rs']
        today_ref = history[-11]['median_rs']
        today_mom = (today_rs - today_ref) / 10.0
        today_q = quadrant_of(today_rs, today_mom)

        # ~5 trading days ago quadrant (~1 week)
        prev_idx = -6 if len(history) >= 16 else -1
        prev_rs = history[prev_idx]['median_rs']
        prev_ref_idx = prev_idx - 10
        if abs(prev_ref_idx) > len(history):
            continue
        prev_ref = history[prev_ref_idx]['median_rs']
        prev_mom = (prev_rs - prev_ref) / 10.0
        prev_q = quadrant_of(prev_rs, prev_mom)

        if today_q != prev_q:
            flips.append({
                'sector': sector,
                'from': prev_q,
                'to': today_q,
                'rs_change': round(today_rs - prev_rs, 2),
                'mom_change': round(today_mom - prev_mom, 3),
            })
    flips.sort(key=lambda f: f['sector'])
    return flips


def fetch_breadth_summary():
    """Last 7 trading days of view_daily_breadth — direct .limit() query
    (NOT paginated_fetch, which was overriding limit with .range)."""
    try:
        resp = supabase.table('view_daily_breadth').select('*') \
            .order('snapshot_date', desc=True).limit(7).execute()
        return resp.data or []
    except Exception as e:
        print(f"    fetch_breadth_summary error: {e}")
        return []


def fetch_top_movers():
    """Top rank surges + vol breakouts on the latest snapshot date."""
    latest_resp = supabase.table('daily_stock_snapshots').select('snapshot_date')\
        .order('snapshot_date', desc=True).limit(1).execute()
    if not latest_resp.data:
        return None
    latest_date = latest_resp.data[0]['snapshot_date']

    snaps = paginated_fetch(
        'daily_stock_snapshots',
        'ticker, rank_slope, vol_ratio, rs_rank, alkalyme_rs, close',
        lambda q: q.eq('snapshot_date', latest_date),
        max_rows=600,
    )
    snaps = [s for s in snaps if s.get('rank_slope') is not None or s.get('vol_ratio') is not None]

    surges = sorted(
        [s for s in snaps if s.get('rank_slope') is not None],
        key=lambda r: r['rank_slope'],
    )[:10]
    declines = sorted(
        [s for s in snaps if s.get('rank_slope') is not None],
        key=lambda r: -r['rank_slope'],
    )[:5]
    vol_bo = sorted(
        [s for s in snaps if s.get('vol_ratio') is not None and s.get('vol_ratio') >= 1.5],
        key=lambda r: -r['vol_ratio'],
    )[:10]

    tickers = list({r['ticker'] for r in surges + declines + vol_bo})
    sector_map = {}
    if tickers:
        sm = supabase.table('indian_stock_sectors')\
            .select('ticker, sector, company_name').in_('ticker', tickers).execute()
        for r in (sm.data or []):
            sector_map[r['ticker']] = r

    def annotate(rows):
        return [{
            **r,
            'sector': sector_map.get(r['ticker'], {}).get('sector', '—'),
            'name': (sector_map.get(r['ticker'], {}).get('company_name') or '')[:32],
        } for r in rows]

    return {
        'date': latest_date,
        'rank_surges': annotate(surges),
        'rank_declines': annotate(declines),
        'vol_breakouts': annotate(vol_bo),
    }


_MONTHS_3 = {'Jan':1,'Feb':2,'Mar':3,'Apr':4,'May':5,'Jun':6,
             'Jul':7,'Aug':8,'Sep':9,'Oct':10,'Nov':11,'Dec':12}

def parse_period_to_date(period_str):
    """Convert 'Mar 2026', 'Mar-2026', 'Mar26', etc. to a date (last day of the
    month) for comparison. Returns None if unparseable."""
    if not period_str:
        return None
    s = str(period_str).strip().replace('-', ' ').replace('/', ' ')
    # Try "Mar 2026" / "March 2026" pattern
    m = re.match(r'^([A-Za-z]{3,9})\s*(\d{2,4})', s)
    if not m:
        return None
    month_token = m.group(1)[:3].title()
    yr_token    = m.group(2)
    month_num   = _MONTHS_3.get(month_token)
    if not month_num:
        return None
    try:
        year = int(yr_token)
        if year < 100:
            year += 2000
    except Exception:
        return None
    from calendar import monthrange
    return date(year, month_num, monthrange(year, month_num)[1])


def fetch_returns_for_tickers(tickers, lookback_calendar_days=14):
    """For each ticker, return latest close + 1-day return + 5-day return
    from daily_stock_snapshots. Returns {ticker: {'close', 'date', 'ret_1d', 'ret_5d'}}."""
    if not tickers:
        return {}
    floor = (date.today() - timedelta(days=lookback_calendar_days)).isoformat()
    rows = paginated_fetch(
        'daily_stock_snapshots',
        'ticker, snapshot_date, close',
        lambda q: q.in_('ticker', tickers).gte('snapshot_date', floor)
                   .order('snapshot_date', desc=True),
    )
    by_ticker = defaultdict(list)
    for r in rows:
        by_ticker[r['ticker']].append(r)

    out = {}
    for ticker, hist in by_ticker.items():
        hist.sort(key=lambda r: r['snapshot_date'], reverse=True)
        if not hist:
            continue
        latest = hist[0].get('close')
        if not latest:
            continue
        ret_1d = None
        if len(hist) >= 2 and hist[1].get('close'):
            ret_1d = (latest - hist[1]['close']) / hist[1]['close'] * 100
        ret_5d = None
        if len(hist) >= 6 and hist[5].get('close'):
            ret_5d = (latest - hist[5]['close']) / hist[5]['close'] * 100
        out[ticker] = {
            'close':    latest,
            'date':     hist[0]['snapshot_date'],
            'ret_1d':   round(ret_1d, 2) if ret_1d is not None else None,
            'ret_5d':   round(ret_5d, 2) if ret_5d is not None else None,
        }
    return out


def fetch_earnings_season():
    """Sector-aggregate medians + top beats/misses from quarterly_financials.
    CRITICAL: only includes stocks whose latest reported quarter matches the
    current results season — otherwise stocks with stale (Dec quarter) data
    contaminate the May aggregate.

    Sectors come from indian_stock_sectors (NSE taxonomy), NOT stock_fundamentals.sector
    (which is Yahoo Finance taxonomy and mostly NULL for non-yfinance-fetched stocks)."""
    rows = paginated_fetch(
        'stock_fundamentals',
        'ticker, quarterly_financials, earnings_growth_yoy, revenue_growth_yoy, market_cap_cr',
    )
    sector_lookup = {}
    try:
        sec_rows = paginated_fetch('indian_stock_sectors', 'ticker, sector')
        for r in sec_rows:
            sector_lookup[r['ticker']] = r.get('sector') or 'Unknown'
    except Exception as e:
        print(f"    sector_lookup error: {e}")

    # Pass 1: parse all stocks, collect latest period per stock
    parsed = []
    for r in rows:
        qf = r.get('quarterly_financials')
        if not qf:
            continue
        if isinstance(qf, str):
            try:
                qf = json.loads(qf)
            except Exception:
                continue
        if not isinstance(qf, list) or len(qf) < 5:
            continue
        period_str  = qf[0].get('period') or qf[0].get('quarter') or ''
        period_date = parse_period_to_date(period_str)
        if period_date is None:
            continue
        parsed.append({
            'ticker':         r['ticker'],
            'qf':             qf,
            'period_str':     period_str,
            'period_date':    period_date,
            'market_cap_cr':  r.get('market_cap_cr'),
        })

    if not parsed:
        return {'sector_medians': [], 'top_beats': [], 'top_misses': [], 'season_period': None, 'stocks_in_season': 0, 'stocks_with_stale_data': 0}

    # The current season = the most recent quarter present in the data
    latest_period_date = max(p['period_date'] for p in parsed)
    latest_period_str  = next(p['period_str'] for p in parsed if p['period_date'] == latest_period_date)

    # Filter to ONLY stocks whose latest reported quarter is the current season
    in_season = [p for p in parsed if p['period_date'] == latest_period_date]
    stale_count = len(parsed) - len(in_season)

    by_sector = defaultdict(list)
    individual = []

    for p in in_season:
        qf = p['qf']
        cur = qf[0].get('eps')
        ago = qf[4].get('eps')
        if cur is None or ago is None or ago == 0:
            continue
        try:
            eps_yoy = (cur - ago) / abs(ago) * 100
        except Exception:
            continue
        if abs(eps_yoy) > 1000:
            continue
        rev_cur = qf[0].get('revenue_cr') or qf[0].get('revenue')
        rev_ago = qf[4].get('revenue_cr') or qf[4].get('revenue')
        rev_yoy = None
        if rev_cur and rev_ago and rev_ago != 0:
            try:
                rev_yoy = (rev_cur - rev_ago) / abs(rev_ago) * 100
            except Exception:
                pass

        sector = sector_lookup.get(p['ticker'], 'Unknown')
        item = {
            'ticker':         p['ticker'],
            'sector':         sector,
            'eps_yoy':        round(eps_yoy, 1),
            'rev_yoy':        round(rev_yoy, 1) if rev_yoy is not None else None,
            'period':         p['period_str'],
            'market_cap_cr':  p.get('market_cap_cr'),
        }
        individual.append(item)
        by_sector[sector].append(eps_yoy)

    sector_aggs = []
    for sector, vals in by_sector.items():
        if len(vals) < 3:
            continue
        vals_sorted = sorted(vals)
        median = vals_sorted[len(vals_sorted) // 2]
        sector_aggs.append({'sector': sector, 'median_eps_yoy': round(median, 1), 'n': len(vals)})
    sector_aggs.sort(key=lambda r: -r['median_eps_yoy'])

    individual.sort(key=lambda r: -r['eps_yoy'])

    # ── Bellwethers: top 25 in-season stocks by market cap (regardless of EPS rank) ──
    # These are sector-moving large caps — SBI, Titan, Reliance, etc. Their post-
    # earnings price reaction often matters more than their EPS YoY %.
    with_mcap = [r for r in individual if r.get('market_cap_cr') and r['market_cap_cr'] > 0]
    with_mcap.sort(key=lambda r: -r['market_cap_cr'])
    bellwether_candidates = with_mcap[:25]

    # Fetch 1d / 5d returns for these bellwethers
    bw_returns = fetch_returns_for_tickers([b['ticker'] for b in bellwether_candidates])
    bellwethers = []
    for b in bellwether_candidates:
        ret = bw_returns.get(b['ticker'], {})
        bellwethers.append({
            **b,
            'ret_1d':   ret.get('ret_1d'),
            'ret_5d':   ret.get('ret_5d'),
            'close':    ret.get('close'),
        })

    return {
        'sector_medians':         sector_aggs,
        'top_beats':              individual[:10],
        'top_misses':             individual[-5:][::-1],
        'bellwethers':            bellwethers,
        'season_period':          latest_period_str,
        'stocks_in_season':       len(individual),
        'stocks_with_stale_data': stale_count,
    }


def fetch_signals_summary():
    """Entry signals fired in last 14 days, split into this-week vs last-week.
    Counts UNIQUE TICKERS per indicator (not signal-rows — a stock firing the
    same indicator on multiple days is one stock). Also identifies stocks
    triggering 2+ DIFFERENT indicator types this week."""
    today = date.today()
    week_ago      = (today - timedelta(days=7)).isoformat()
    two_weeks_ago = (today - timedelta(days=14)).isoformat()

    rows = paginated_fetch(
        'entry_signals',
        'ticker, stock_name, signal_type, signal_strength, alert_date, details',
        lambda q: q.gte('alert_date', two_weeks_ago).order('alert_date', desc=True),
    )

    # Split into this-week / prior-week
    this_week, prior_week = [], []
    for r in rows:
        if r['alert_date'] >= week_ago:
            this_week.append(r)
        else:
            prior_week.append(r)

    # Unique tickers per signal_type (this week vs prior week)
    this_by_type  = defaultdict(set)
    prior_by_type = defaultdict(set)
    examples_by_type = defaultdict(list)
    for r in this_week:
        this_by_type[r['signal_type']].add(r['ticker'])
        examples_by_type[r['signal_type']].append(r)
    for r in prior_week:
        prior_by_type[r['signal_type']].add(r['ticker'])

    # Build per-indicator summary with unique counts + WoW delta + top tickers
    all_types = set(this_by_type.keys()) | set(prior_by_type.keys())
    summary = []
    for stype in all_types:
        this_set  = this_by_type.get(stype, set())
        prior_set = prior_by_type.get(stype, set())

        # Top examples for this week, deduped by ticker, ranked by RS
        seen, ranked = set(), []
        for r in examples_by_type.get(stype, []):
            if r['ticker'] in seen:
                continue
            seen.add(r['ticker'])
            details = r.get('details') if isinstance(r.get('details'), dict) else {}
            rank = details.get('rs_rank') if isinstance(details, dict) else None
            ranked.append((rank or 9999, r))
        ranked.sort(key=lambda x: x[0])

        summary.append({
            'type':             stype,
            'unique_this_week': len(this_set),
            'unique_prior':     len(prior_set),
            'delta':            len(this_set) - len(prior_set),
            'fresh_entries':    len(this_set - prior_set),  # stocks new this week
            'dropped':          len(prior_set - this_set),  # stocks gone from list
            'top_examples':     [r for _, r in ranked[:6]],
        })
    summary.sort(key=lambda r: -r['unique_this_week'])

    # Stocks firing 2+ DIFFERENT indicator types this week (multi-signal cluster)
    ticker_to_types = defaultdict(set)
    for r in this_week:
        ticker_to_types[r['ticker']].add(r['signal_type'])
    multi = [(t, sorted(types)) for t, types in ticker_to_types.items() if len(types) >= 2]
    multi.sort(key=lambda x: -len(x[1]))

    return {
        'by_type':        summary,
        'multi_signal':   multi[:15],
        'total_unique_this_week': len({r['ticker'] for r in this_week}),
    }


# =============================================================================
# NEWS FETCHER (Google News RSS)
# =============================================================================

# Catalyst keywords — a headline mentioning any of these names an actual
# market-moving cause, so it is prioritised over generic "why did the market
# crash today" clickbait when the feed is trimmed to the context cap.
CATALYST_KEYWORDS = re.compile(
    r'\b(duty|tariff|tax|gst|budget|fiscal|excise|subsidy|import|export|'
    r'petrol|diesel|fuel|crude|oil|brent|rbi|repo|inflation|cpi|gdp|rupee|'
    r'fed|fii|dii|sebi|austerity|earnings|results|stake|merger|acquisition|'
    r'policy|sanction|tension|war|deal)\b',
    re.IGNORECASE,
)


def fetch_google_news():
    queries = [
        '"Nifty 50"',
        '"Sensex"',
        '"Indian stock market"',
        '"RBI" India',
        '"Indian economy" markets',
        '"IPO" India',
        # Policy / fiscal / commodity catalysts. The index-focused queries
        # above surface "the market fell" headlines but rarely the WHY —
        # duty changes, fuel-price revisions and govt measures. These do.
        'India "import duty" OR "customs duty" OR tariff',
        'India petrol diesel price',
        'India government fiscal OR tax OR GST OR budget',
        'India gold silver duty OR price',
        'India "PM Modi" economy OR austerity OR policy',
    ]
    all_items = []
    for q in queries:
        url = (
            f'https://news.google.com/rss/search?'
            f'q={urllib.parse.quote(q)}+when:7d&hl=en-IN&gl=IN&ceid=IN:en'
        )
        try:
            feed = feedparser.parse(url)
            for entry in feed.entries[:15]:
                title = html.unescape(entry.get('title', '')).strip()
                if not title:
                    continue
                source = ''
                if hasattr(entry, 'source') and entry.source:
                    source = getattr(entry.source, 'title', '') or ''
                all_items.append({
                    'title': title,
                    'source': source,
                    'published': entry.get('published', ''),
                    'link': entry.get('link', ''),
                })
        except Exception as e:
            print(f"    Google News error for {q}: {e}")
            continue

    seen, unique = set(), []
    for item in all_items:
        key = re.sub(r'\W+', '', item['title'].lower())[:60]
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)

    # Stable sort: headlines naming a concrete catalyst float to the top so
    # they survive the cap; generic clickbait sinks but is not discarded.
    unique.sort(key=lambda x: 0 if CATALYST_KEYWORDS.search(x['title']) else 1)
    return unique[:40]


# =============================================================================
# NEWS EVENT EXTRACTION (Claude Haiku pre-pass)
# =============================================================================

NEWS_EXTRACT_PROMPT = """You are a financial news analyst. You will receive news headlines about
Indian markets from the past 7 days. Extract the distinct MARKET-MOVING
EVENTS behind them.

Rules:
- Merge every headline about the same underlying event into ONE row.
- IGNORE pure clickbait that names no catalyst ("Why did the market crash
  today?", "Sensex crashes 1,300 points", "5 key factors", trading-plan and
  price-prediction headlines). They carry no information.
- Pay special attention to DOMESTIC POLICY and FISCAL catalysts: import/export
  duty or tariff changes, fuel/excise price revisions, tax/GST changes,
  government austerity or spending measures, regulatory (RBI/SEBI) actions.
  These are often the real driver of a move and the easiest to miss.
- For each event provide: a one-line description; a TYPE (exactly one of:
  Fiscal/Policy, Monetary/RBI, Commodity, Geopolitical, Corporate/Earnings,
  Global, Other); WHEN (exactly one of: "Occurred this week",
  "Upcoming (dated)", "Unclear/ongoing"); the AFFECTED sectors or assets; and
  the likely IMPACT on Indian equities (Bullish / Bearish / Mixed).
- Do NOT invent events no headline supports. If a single vague headline only
  hints at something, still include it but end the description with
  "(single vague headline — unconfirmed)".
- If the headlines contain no clear market-moving event, output exactly:
  "No distinct market-moving events identified."

Output ONLY a Markdown table with these columns:
| Event | Type | When | Affected | Impact |
No preamble, no commentary before or after the table."""


def extract_news_events(news):
    """Haiku pre-pass: digest raw headlines into a structured event table.

    Returns a Markdown table string, or None if extraction is unavailable
    (caller falls back to the raw headline list)."""
    if not news:
        return None
    lines = []
    for i, n in enumerate(news, 1):
        src = f" ({n['source']})" if n.get('source') else ''
        lines.append(f"{i}. {n['title']}{src}")
    headline_block = '\n'.join(lines)
    try:
        resp = claude.messages.create(
            model=HAIKU_MODEL,
            max_tokens=1500,
            system=NEWS_EXTRACT_PROMPT,
            messages=[{
                'role': 'user',
                'content': f"Headlines (past 7 days):\n\n{headline_block}",
            }],
        )
        digest = resp.content[0].text.strip()
        return digest or None
    except Exception as e:
        print(f"    ⚠️  News event extraction failed: {e}")
        return None


# =============================================================================
# CONTEXT BUILDER
# =============================================================================

def fmt_pct(v):
    if v is None:
        return '—'
    return f"{'+' if v >= 0 else ''}{v:.1f}%"


def fmt_num(v, decimals=2):
    if v is None:
        return '—'
    return f"{v:,.{decimals}f}"


def build_context(macro, indexes, nifty_emas, sectors, flips, breadth, movers, earnings, signals, news, news_digest=None):
    L = []
    today = date.today().strftime('%Y-%m-%d')
    L.append(f"# Weekly Market Context — {today}\n")

    # ── Daily Indian index closes — AUTHORITATIVE ground truth for date↔value ──
    L.append("## This Week's Indian Index Closes (AUTHORITATIVE — use these for any day-of-week claims)\n")
    if indexes and indexes.get('days'):
        L.append("| Date | Day | Nifty 50 | Sensex | NiftyBees |")
        L.append("|---|---|---|---|---|")
        for d in indexes['days']:
            L.append(
                f"| {d['date']} | {d['day']} | "
                f"{fmt_num(d.get('nifty'))} | {fmt_num(d.get('sensex'))} | "
                f"{fmt_num(d.get('niftybees'))} |"
            )
        if 'nifty_pct' in indexes:
            L.append(
                f"\n**Nifty 50 week summary** ({indexes['date_first']} → {indexes['date_last']}): "
                f"{fmt_num(indexes['nifty_first'])} → {fmt_num(indexes['nifty_close'])} "
                f"({fmt_pct(indexes['nifty_pct'])} for the week)."
            )
        if 'sensex_pct' in indexes:
            L.append(
                f"**Sensex week summary**: {fmt_num(indexes['sensex_first'])} → "
                f"{fmt_num(indexes['sensex_close'])} ({fmt_pct(indexes['sensex_pct'])} for the week)."
            )
    else:
        L.append("_No detailed_snapshots data available._")
    L.append("")

    # ── Nifty EMA trend levels — AUTHORITATIVE for "Nifty vs MA" claims ──
    L.append("## Nifty 50 Trend Levels (AUTHORITATIVE — use for any Nifty-vs-MA claim)\n")
    if nifty_emas:
        latest_close = None
        if indexes and indexes.get('days'):
            # Prefer the most recent value from detailed_snapshots (intraday-accurate)
            latest_close = indexes['days'][-1].get('nifty')
        if latest_close is None:
            latest_close = nifty_emas['last_close']
        e20  = nifty_emas.get('ema_20')
        e50  = nifty_emas.get('ema_50')
        e200 = nifty_emas.get('ema_200')
        def status(close, ema):
            if ema is None or close is None: return '—'
            diff_pct = (close - ema) / ema * 100
            return f"{'ABOVE' if close > ema else 'BELOW'} ({fmt_pct(diff_pct)})"
        L.append(f"- **Latest Nifty close**: {fmt_num(latest_close)}")
        L.append(f"- **EMA 20**: {fmt_num(e20)} — Nifty is {status(latest_close, e20)}")
        L.append(f"- **EMA 50**: {fmt_num(e50)} — Nifty is {status(latest_close, e50)}")
        L.append(f"- **EMA 200**: {fmt_num(e200)} — Nifty is {status(latest_close, e200)}")
        L.append(f"\n_(Computed from {nifty_emas.get('days_used', 0)} days of market_ratios history.)_")
    else:
        L.append("_Insufficient history to compute EMAs._")
    L.append("")

    # Macro
    L.append("## Macro Snapshot — global / commodities / valuations\n")
    if macro and macro.get('items'):
        L.append(f"_(Latest: {macro['latest_date']} · week-ago ref: {macro['week_ago_date']})_\n")
        for item in macro['items']:
            line = f"- **{item['label']}**: {fmt_num(item['current'])}"
            if 'change_pct' in item:
                line += f" ({fmt_pct(item['change_pct'])} week)"
            elif 'change_abs' in item:
                ab = item['change_abs']
                line += f" ({'+' if ab >= 0 else ''}{ab:.2f} week)"
            L.append(line)
    else:
        L.append("_No macro data available._")
    L.append("")

    # Sector summary
    L.append("## Sector Performance — 5-day median return\n")
    if sectors:
        L.append("**TOP 5:**")
        L.append("| Sector | 5D | 1M | 3M | 6M | YTD | 1Y |")
        L.append("|---|---|---|---|---|---|---|")
        for s in sectors.get('top_5', []):
            L.append(
                f"| {s['sector']} | {fmt_pct(s.get('med_5d'))} | "
                f"{fmt_pct(s.get('med_1m'))} | {fmt_pct(s.get('med_3m'))} | "
                f"{fmt_pct(s.get('med_6m'))} | {fmt_pct(s.get('med_ytd'))} | "
                f"{fmt_pct(s.get('med_1y'))} |"
            )
        L.append("")
        L.append("**BOTTOM 5:**")
        L.append("| Sector | 5D | 1M | 3M | 6M | YTD | 1Y |")
        L.append("|---|---|---|---|---|---|---|")
        for s in sectors.get('bottom_5', []):
            L.append(
                f"| {s['sector']} | {fmt_pct(s.get('med_5d'))} | "
                f"{fmt_pct(s.get('med_1m'))} | {fmt_pct(s.get('med_3m'))} | "
                f"{fmt_pct(s.get('med_6m'))} | {fmt_pct(s.get('med_ytd'))} | "
                f"{fmt_pct(s.get('med_1y'))} |"
            )
    L.append("")

    # Quadrant flips
    L.append("## Sector Quadrant Flips (this week)\n")
    if flips:
        for f in flips:
            L.append(
                f"- **{f['sector']}**: {f['from']} → {f['to']} "
                f"(RS Δ {f['rs_change']:+.1f}, momentum Δ {f['mom_change']:+.2f})"
            )
    else:
        L.append("_No sector quadrant flips this week._")
    L.append("")

    # Breadth trajectory
    L.append("## Breadth Trajectory (last 7 trading days)\n")
    if breadth:
        L.append("| Date | Adv | Dec | Unc | New 52W H | New 52W L | % > 200 EMA |")
        L.append("|---|---|---|---|---|---|---|")
        for b in reversed(breadth):  # chronological order
            total = b.get('total_stocks') or 1
            ema_pct = (b.get('above_ema_200') or 0) / total * 100 if total else 0
            L.append(
                f"| {b['snapshot_date']} | {b.get('advancing','—')} | "
                f"{b.get('declining','—')} | {b.get('unchanged','—')} | "
                f"{b.get('new_highs_52w','—')} | {b.get('new_lows_52w','—')} | "
                f"{ema_pct:.0f}% |"
            )
    L.append("")

    # Top movers
    if movers:
        L.append(f"## Top 10 Rank Surges  ({movers.get('date','')})\n")
        L.append("| Ticker | Name | Sector | Slope (pos/day) | Now Ranked | RS |")
        L.append("|---|---|---|---|---|---|")
        for r in movers.get('rank_surges', []):
            L.append(
                f"| {r['ticker'].replace('.NS','')} | {r.get('name','')} | "
                f"{r.get('sector','—')} | {r.get('rank_slope', 0):.1f} | "
                f"#{r.get('rs_rank') or '—'} | "
                f"{(r.get('alkalyme_rs') or 0):.1f} |"
            )
        L.append("")
        L.append("**Top 5 Rank Decliners (worst rank slope):**")
        L.append("| Ticker | Sector | Slope | Now Ranked |")
        L.append("|---|---|---|---|")
        for r in movers.get('rank_declines', []):
            L.append(
                f"| {r['ticker'].replace('.NS','')} | {r.get('sector','—')} | "
                f"{r.get('rank_slope', 0):.1f} | #{r.get('rs_rank') or '—'} |"
            )
        L.append("")
        L.append("**Top 10 Volume Breakouts (vol_ratio ≥ 1.5x):**")
        L.append("| Ticker | Name | Sector | Vol Ratio | RS Rank |")
        L.append("|---|---|---|---|---|")
        for r in movers.get('vol_breakouts', []):
            L.append(
                f"| {r['ticker'].replace('.NS','')} | {r.get('name','')} | "
                f"{r.get('sector','—')} | {r.get('vol_ratio', 0):.1f}x | "
                f"#{r.get('rs_rank') or '—'} |"
            )
    L.append("")

    # Earnings — filtered to current results-season quarter only
    if earnings and earnings.get('season_period'):
        L.append(f"## Earnings — {earnings['season_period']} Results Season Only\n")
        L.append(
            f"_Filtered to stocks whose latest reported quarter is **{earnings['season_period']}**. "
            f"{earnings['stocks_in_season']} stocks in season; "
            f"{earnings['stocks_with_stale_data']} stocks excluded (haven't reported the latest quarter yet — "
            f"their data is from a prior quarter and would distort the season aggregate)._\n"
        )
        L.append("**Sector medians (YoY EPS):**")
        L.append("| Sector | Median EPS YoY | # of stocks |")
        L.append("|---|---|---|")
        for s in earnings.get('sector_medians', [])[:18]:
            L.append(f"| {s['sector']} | {s['median_eps_yoy']:+.1f}% | {s['n']} |")
        L.append("")

        L.append(f"**Top 10 EPS Beats this season ({earnings['season_period']}):**")
        L.append("| Ticker | Sector | EPS YoY | Rev YoY |")
        L.append("|---|---|---|---|")
        for r in earnings.get('top_beats', []):
            L.append(
                f"| {r['ticker'].replace('.NS','')} | {r['sector']} | "
                f"{r['eps_yoy']:+.0f}% | {fmt_pct(r.get('rev_yoy'))} |"
            )
        L.append("")

        L.append(f"**Top 5 EPS Misses this season ({earnings['season_period']}):**")
        L.append("| Ticker | Sector | EPS YoY | Rev YoY |")
        L.append("|---|---|---|---|")
        for r in earnings.get('top_misses', []):
            L.append(
                f"| {r['ticker'].replace('.NS','')} | {r['sector']} | "
                f"{r['eps_yoy']:+.0f}% | {fmt_pct(r.get('rev_yoy'))} |"
            )
        L.append("")

        # ── BELLWETHER LARGE-CAPS: sector-moving stocks regardless of EPS % rank ──
        if earnings.get('bellwethers'):
            L.append(f"**Bellwether Large-Cap Earnings ({earnings['season_period']}) — top 25 by market cap, with post-earnings price action:**")
            L.append("_These are the stocks whose results move their sector regardless of where they rank by % YoY change. The 1-day and 5-day returns capture market reaction — a stock that fell 6% on a +5% EPS print is telling you something different from one that rallied 8% on the same number._\n")
            L.append("| Ticker | Sector | Mkt Cap (₹ cr) | EPS YoY | Rev YoY | 1-day | 5-day |")
            L.append("|---|---|---|---|---|---|---|")
            for b in earnings['bellwethers']:
                mc = b.get('market_cap_cr') or 0
                mc_fmt = f"{mc:,.0f}"
                L.append(
                    f"| {b['ticker'].replace('.NS','')} | {b.get('sector','—')} | "
                    f"{mc_fmt} | {b['eps_yoy']:+.0f}% | {fmt_pct(b.get('rev_yoy'))} | "
                    f"{fmt_pct(b.get('ret_1d'))} | {fmt_pct(b.get('ret_5d'))} |"
                )
    L.append("")

    # Signals — UNIQUE STOCK COUNTS this week vs prior week
    L.append("## Technical Signals — Unique Stocks This Week vs Prior Week\n")
    if signals and signals.get('by_type'):
        L.append(
            f"_All counts are **unique stocks** (a stock firing the same indicator on multiple "
            f"days = 1 stock). Total unique stocks across all indicators this week: "
            f"**{signals.get('total_unique_this_week', 0)}**._\n"
        )
        L.append("| Indicator | This Week (unique) | Prior Week | Δ | Fresh entries | Top examples |")
        L.append("|---|---|---|---|---|---|")
        for s in signals['by_type']:
            tickers = ', '.join(ex['ticker'].replace('.NS', '') for ex in s.get('top_examples', [])[:5])
            delta = s['delta']
            delta_str = f"{'+' if delta > 0 else ''}{delta}"
            L.append(
                f"| {s['type']} | {s['unique_this_week']} | {s['unique_prior']} | "
                f"{delta_str} | {s['fresh_entries']} | {tickers} |"
            )
        L.append("")

        if signals.get('multi_signal'):
            L.append("**Stocks triggering 2+ different indicator types this week (cluster signals):**")
            for ticker, types in signals['multi_signal'][:10]:
                clean = ticker.replace('.NS', '')
                L.append(f"- **{clean}**: {len(types)} signals — {', '.join(types)}")
        else:
            L.append("_No stocks triggered multiple indicator types this week._")
    else:
        L.append("_No entry signals fired in the last 7 days._")
    L.append("")

    # Market-moving events digest (Haiku pre-pass over the raw headlines)
    L.append("## Market-Moving Events This Week (categorised from news headlines)\n")
    if news_digest:
        L.append("_A news-analysis pass grouped this week's headlines into distinct "
                 "events. This is the PRIMARY source for WHY the market moved — "
                 "especially domestic policy / fiscal catalysts. The data tables "
                 "above are the source for the numbers; this section is the source "
                 "for the cause._\n")
        L.append(news_digest)
    else:
        L.append("_Event extraction unavailable this run — rely on the raw "
                 "headlines below for cause / narrative._")
    L.append("")

    # News (raw headlines — kept for transparency and as a fallback)
    L.append("## News Headlines (last 7 days, deduped — raw feed)\n")
    if news:
        for n in news[:30]:
            src = f" — _{n['source']}_" if n.get('source') else ''
            L.append(f"- {n['title']}{src}")
    else:
        L.append("_No news fetched._")

    return '\n'.join(L)


# =============================================================================
# CLAUDE PROMPTS
# =============================================================================

LONG_FORM_PROMPT = """You are a market intelligence writer producing a weekly report on Indian
equity markets. Your audience is mixed: some are seasoned investors comfortable
with technical jargon (RRG quadrants, Blue Zone, sector rotation), others are
starting out and need plain-English explanations.

# Audience rule (strict)
On the FIRST use of any technical term, pair it with a one-line plain-English
explanation in parentheses. Examples:
- "Blue Zone (the dashboard's institutional-accumulation watchlist) jumped from 104 to 118 stocks."
- "Bull Flag breakout (chart pattern: sharp price gain followed by orderly consolidation, then a breakout above the consolidation range)..."
- "Relative Strength rank (where each sector ranks vs all others on momentum strength — lower number = stronger)..."

# Style
- Conversational but data-driven. Cite specific numbers everywhere.
- Connect dots across data sources — link macro moves -> sector rotation -> specific earnings -> individual stock action.
- Be honest when data is mixed or contradicts itself.
- No stock recommendations. Use "worth watching", "early signs of", "sector showing strength" — never "buy" or "sell".
- Indian formatting: use rupee sign for currency, lakhs/crores where natural.
- Reference days by name ("Wednesday's close", "by Friday").

# Tone
- Confident but humble.
- A little personality OK ("textbook stealth-strength day", "the rotation is happening in real time"). Not gimmicky.
- Educational asides for beginners woven in naturally — never preachy.

# CRITICAL — accuracy rules
- For day-of-week claims (e.g. "Tuesday's surge"), use ONLY the dates from the
  "This Week's Indian Index Closes" table at the top of the context. That table
  is the AUTHORITATIVE source for which date was Mon/Tue/Wed/Thu/Fri.
- For Nifty/Sensex closes, use ONLY values from that same table.
- For "Nifty above/below 20-day / 50-day / 200-day moving average" claims, use
  ONLY the "Nifty 50 Trend Levels" section. It tells you exactly which EMAs
  Nifty is above and which it's below. NEVER assume — read the explicit status.
- For earnings, ONLY discuss stocks/sectors from the filtered current-season
  data. The earnings section explicitly states the quarter being analysed.
  Do NOT cite numbers from stocks excluded as "stale data" (those are from a
  prior quarter and would mislead readers).
- For technical signals, use UNIQUE STOCK COUNTS (not "X signals fired").
  The signals table provides this-week vs prior-week unique counts and the
  delta. Frame the conversation as "the Blue Zone universe expanded from N to M
  unique stocks this week" — NOT "M signals were fired."
- TWO separate rules govern the news data — do not conflate them:
  (a) FACTS: a headline may state a date, index level, or % inaccurately.
      NEVER take a date, price, EMA or percentage from a headline or from the
      events digest — those come ONLY from the data tables above.
  (b) CAUSE: the "Market-Moving Events This Week" section (and the raw
      headlines beneath it) ARE your primary, authoritative source for WHY
      the market moved. The data tables tell you WHAT moved; the news tells
      you WHY. Use it actively — never leave the "why" of a move vague.
- When explaining why an index or sector moved, consult the "Market-Moving
  Events" digest FIRST. Domestic policy and fiscal catalysts — import/export
  duty or tariff changes, fuel/excise price revisions, tax/GST changes,
  government austerity or spending measures, RBI/SEBI actions — are frequently
  the real driver and are easy to under-weight. Do NOT default to the largest
  quantified macro number (e.g. a Brent crude move) as the explanation when a
  policy event in the digest is a more direct cause.
- DIVERGENCE TELL: when a domestic sector or ETF moves OPPOSITE to its global
  counterpart — e.g. Indian gold/silver ETFs rising while global gold (USD)
  falls — that divergence almost always signals a DOMESTIC policy or tax cause
  (an import-duty hike lifts the rupee landed price regardless of the global
  price). Find the cause in the events digest and name it explicitly — do not
  label it a generic "safe-haven bid".
- If a data point isn't in the context, DON'T invent it.

# Structure (~1400-1800 words, 9 sections)

## 1. What Happened This Week
3-4 sentences. Headline week % move (use the index-closes table) + the dominant
theme (rotation, breadth shift, macro event, single-stock blowup, etc.) AND the
dominant CAUSE — name the specific catalyst(s) from the "Market-Moving Events"
digest that drove the week. If a domestic policy or fiscal event is in the
digest, it almost certainly belongs in this opening paragraph.

## 2. The Trend Picture
Nifty's position vs its 20 / 50 / 200-day EMAs (Exponential Moving Average —
a smoothed average of recent closing prices that traders use as trend
reference). Use the EXPLICIT "Nifty 50 Trend Levels" section in the context
to state whether Nifty is ABOVE or BELOW each one. Don't guess — the section
tells you. Then characterise the trend (broadening / narrowing / mixed). Use
breadth numbers: % stocks above EMA200, daily A/D, fresh 52W highs vs lows.
Reference the daily breadth table (last 7 days) — call out the day with the
strongest A/D split.

## 3. Sector Rotation — Where Money Is Flowing
Top 3 / bottom 3 sectors by weekly return. Highlight any quadrant flips
(Leading <-> Weakening, etc.). Link each notable sector move to a cause from
the "Market-Moving Events" digest or the macro data — e.g. "Brent fell 8%, so
auto parts and aviation rose" or "the gold/silver import-duty hike lifted ETF
prices". Apply the divergence tell: a domestic ETF/sector moving against its
global counterpart points to a policy cause, not a flow-of-funds story.

## 4. Earnings Pulse
The earnings section is filtered to ONE specific quarter — state which
quarter at the top (read it from the section header). Mention how many
stocks have reported and how many haven't yet. Then cover THREE things:

**(a) Sector aggregate**: which sectors are reporting strong (median EPS YoY
positive), which weak. Identify sector patterns ("FMCG quarter is broadly
strong — 5 of 8 reported names beat with double-digit EPS").

**(b) Top-of-the-list extremes** (1-2 lines): biggest EPS beats and misses
with the YoY number, but ONLY mention these as colour — do NOT lean on them
as sector signals (a +700% EPS print is often a low-base effect on a small
stock and doesn't move a sector).

**(c) BELLWETHER LARGE-CAPS — REQUIRED COVERAGE**: this is the most
important earnings sub-section. Use the "Bellwether Large-Cap Earnings"
table. These are sector-moving stocks (SBI, Reliance, Titan, HDFC Bank, TCS,
etc.) — the market's reaction to *their* numbers tells you what's actually
happening to a sector. Pick 3-5 bellwethers that had the most
information-rich result (either large price move, or surprising
beat/disappointment) and write a 2-3 sentence read on each:
- The print (EPS YoY + Rev YoY).
- The market reaction (1d / 5d return).
- WHY the reaction (if obvious from the numbers — margin compression,
  guidance hint, sector context, etc.). If the reaction doesn't match the
  print (e.g. SBI fell 6% on +5% PAT), that's the most important story —
  call it out and offer the most likely explanation from the data, with
  honest caveat ("the move suggests X, though we'd need the call transcript
  to confirm").
- Sector implication ("read-through for other PSU banks", "suggests
  premium-discretionary spend is holding up", etc.).

Do NOT mention stocks that aren't in the filtered list — they're from a
prior quarter.

## 5. Technical Signals This Week — REQUIRED COVERAGE
Use the "Technical Signals — Unique Stocks This Week vs Prior Week" section.
You MUST mention EVERY indicator type that has a non-zero count, framed as
**unique stock counts** (NOT "X signals fired"). Cite the WoW delta wherever
present — that's the most informative number ("Blue Zone Buy expanded from
102 to 87 unique stocks, a contraction of 15 — the indicator's universe is
shrinking, suggesting the broad accumulation theme is losing breadth.")

For each indicator:
- One-line plain-English explanation on first use.
- Unique stocks this week vs prior week + delta.
- 1-2 example tickers from "top examples" (clean tickers, no .NS suffix).

Indicators (skip any with 0 stocks this week):
- **Blue Zone Buy / Strong** — momentum stocks under accumulation (RSI EMA above thresholds + close above EMA20 + within 10% of 52W high).
- **Golden Cross Buy** — 20-day EMA crossing above 50-day EMA (medium-term bullish trend signal).
- **MACD Buy / Strong** — momentum oscillator turning positive (trend-change signal).
- **Bull Flag Buy** — sharp rally followed by orderly consolidation, then breakout above the consolidation range.
- **VCP Buy** — Volatility Contraction Pattern (Minervini): tightening price range before a breakout.
- **Pullback Bounce** — strong-trend stocks dipping to their EMA20 and recovering.
- **Narrow CPR** — Central Pivot Range setup, often a breakout precursor.

Then a short paragraph on **stocks triggering 2+ DIFFERENT indicator types
this week** (the "multi-signal" callouts in the context). Multi-signal stocks
are higher-conviction setups — name 3-5 of them with which signals they
triggered. e.g., "HFCL fired both Blue Zone Strong and Bull Flag Buy this
week — the kind of cross-indicator confirmation that suggests stronger
follow-through."

End with one sentence on what the signal mix is telling us this week:
expanding (rally broadening), contracting (rally narrowing), or rotating
(certain indicators surging while others fade).

## 6. Notable Stock-Level Action
3-5 stories from the rank-surge + volume-breakout tables. Big rank surges,
single-day vol explosions, sector-cluster events ("two healthcare diagnostics
ripping the same day is not coincidence"). One sentence each.

## 7. Macro & Global
Brent / USDINR / DXY / Gold / VIX week-over-week moves (from Macro Snapshot
section, which is fine for these). Brief US / global context. Then cover the
domestic macro and policy events from the "Market-Moving Events" digest —
RBI/SEBI actions, fiscal measures, duty/tax/fuel-price changes, government
announcements — and their market impact. This section MUST reflect every
Fiscal/Policy and Monetary/RBI event in the digest marked as occurring this
week; do not omit one because it lacks a tidy quantified figure.

## 8. What to Watch Next Week
ONLY mention items that are clearly, verifiably in the FUTURE. The LLM has
no calendar — only news headlines from the past 7 days. Headlines often
discuss aftermath of past events (state election results already announced,
budget already presented, RBI decision already taken) which can be mistaken
for upcoming events. Apply this rule strictly:

- **Earnings due**: ONLY if a headline explicitly flags upcoming corporate
  results (e.g., "Reliance to report Q4 results next Wednesday"). If no such
  headlines exist in the feed, skip this bullet entirely. Do NOT make up a
  generic "watch earnings next week" line.
- **Macro events**: ONLY if a headline, or a "Market-Moving Events" digest row
  whose When column reads "Upcoming (dated)", explicitly flags a future RBI
  meeting, GDP / CPI print, Fed decision, etc. Digest rows marked "Occurred
  this week" are PAST — never list them here. Avoid vague "watch for any new
  RBI signal" — not actionable.
- **Political / geopolitical events**: BE EXTRA CAREFUL. State election
  results, budget announcements, treaty signings often LOOK upcoming in
  headlines but have actually already happened. If a headline says
  "election results impact on markets", the result is in the PAST, not
  upcoming. When unclear, OMIT — better to skip than to misinform readers.
- **Levels to watch on Nifty / Bank Nifty / sector indices**: ALWAYS safe
  to include — these are technical, not event-based. Cite specific levels
  from the index-closes table (e.g., "watch whether Nifty holds above
  Friday's [close]" or "next overhead resistance is the recent high near
  [value]").
- **Continuation themes**: also safe — "Watch whether the auto rotation
  continues into second-tier names" or "Pharma earnings tail to monitor"
  is fine because it's framed as ongoing observation, not specific event.

If you're uncertain whether something is upcoming or past, OMIT it.

## 9. The Bottom Line
3 numbered observations. NOT recommendations — framings like:
- "Watch for FMCG follow-through next week — Tata Consumer + Nestle on Tuesday."
- "Auto parts rotation is fresh — confirmation needed via second-tier names."
Include honest caveats.

# Constraints
- No personal portfolio data — this is a market view.
- If a data section is sparse, say so briefly and move on.
- Length: 1400-1700 words. Quality over length.

# Input
You'll receive a structured context block with: macro snapshot, sector summary, breadth, top movers, earnings season + notable prints, signals fired this week, a categorised "Market-Moving Events" digest, and the raw news headlines.

Output: pure Markdown. Use ## for section headers, **bold** for emphasis on numbers, tables (| a | b |) where they make data scan easier. No frontmatter, no signoff.
"""


SHORT_FORM_PROMPT = """You are a market intelligence writer producing a SHORT weekly Indian equity
markets summary for NEW investors.

# Audience
Starting investors. Pair every technical term with a one-line plain-English
explanation on first use. No assumed knowledge of RRG / Darvas / Blue Zone.

# Style
- Concise, scannable. Short sentences.
- Numbers, but don't over-detail.
- Friendly, conversational.
- Educational where natural.

# Structure (~550-700 words)

## The Week in One Line
Single sentence: Nifty's % move + the dominant theme.

## What's Working
3 best-performing sectors. One line each: name + % move + one-sentence "why" (news, earnings, macro).

## What's Not
3 worst-performing sectors. Same format.

## Notable Events
2-3 things that moved the market this week. Use the "Market-Moving Events"
digest in the context as your source for the WHY — especially domestic policy
and fiscal catalysts (duty/tax changes, fuel-price revisions, government
measures), which matter as much as earnings and global macro. One short
paragraph each, in plain English. Never take a date or price from a headline —
only the cause.

## Charts to Know
1-2 visualizations described in words (no images). Example: "The advance/decline line — which counts how many stocks rose vs fell each day — kept climbing all week. That's a sign more stocks are participating in the rally, not just a few large names."

## What to Watch Next Week
3 bullets:
- Major earnings due
- Key macro events
- Levels to watch on Nifty

# Constraints
- No portfolio data, no stock recommendations.
- Brevity is the point. 550-700 words.
- Avoid jargon without explanation.

# Input
Same context block as the long-form report — including a categorised
"Market-Moving Events" digest. Be selective. The digest is your authoritative
source for WHY sectors and the index moved; the data tables are the source for
the numbers. If a domestic policy / fiscal event is in the digest, it belongs
in this report. When a domestic ETF or sector moves opposite to its global
counterpart (e.g. Indian gold ETFs up while global gold falls), that points to
a domestic policy/tax cause — name it, don't call it a "safe-haven" move.

Output: pure Markdown. ## for section headers, **bold** for emphasis. No tables (keep it linear). No frontmatter, no signoff.
"""


def call_claude(system_prompt, context, max_tokens=4500):
    resp = claude.messages.create(
        model=MODEL,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{'role': 'user', 'content': context}],
    )
    return resp.content[0].text


# =============================================================================
# MARKDOWN -> DOCX RENDERER
# =============================================================================

def add_inline_runs(p, text):
    """Parse inline **bold**, *italic*, `code` into runs."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            r = p.add_run(token[2:-2]); r.bold = True
        elif token.startswith('*'):
            r = p.add_run(token[1:-1]); r.italic = True
        elif token.startswith('`'):
            r = p.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def md_to_docx(md_text, output_path, title=None):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        for r in h.runs:
            r.font.name = 'Calibri'

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Tables (markdown table)
        if (line.startswith('|') and i + 1 < len(lines)
                and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip())):
            header = [c.strip() for c in line.strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            t = doc.add_table(rows=1 + len(rows), cols=len(header))
            try:
                t.style = 'Light Grid Accent 1'
            except KeyError:
                pass
            for j, h in enumerate(header):
                cell = t.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
            for ri, row in enumerate(rows):
                for j in range(len(header)):
                    cell = t.rows[ri + 1].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    val = row[j] if j < len(row) else ''
                    add_inline_runs(p, val)
                    for run in p.runs:
                        run.font.size = Pt(10)
            doc.add_paragraph()
            continue

        # Bullets
        if re.match(r'^[\-\*]\s+', line):
            txt = re.sub(r'^[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_runs(p, txt)
            i += 1
            continue

        # Numbered
        if re.match(r'^\d+\.\s+', line):
            txt = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline_runs(p, txt)
            i += 1
            continue

        # Horizontal rule
        if line in ('---', '***', '___'):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            txt = line[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(txt); r.italic = True
            r.font.color.rgb = None  # default secondary
            i += 1
            continue

        # Regular paragraph (multi-line until blank or block)
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not re.match(
                r'^(#{1,6}\s+|\||\d+\.\s+|[\-\*]\s+|---|\*\*\*|___|>\s+)', lines[j]):
            para_lines.append(lines[j].rstrip())
            j += 1
        text = ' '.join(para_lines)
        p = doc.add_paragraph()
        add_inline_runs(p, text)
        i = j

    doc.save(output_path)


# =============================================================================
# MAIN
# =============================================================================

def main():
    print("=" * 60)
    print("Weekly Market Intelligence Report")
    print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)

    print("\n📊 Assembling Supabase data...")
    indexes    = fetch_daily_index_closes();   print(f"  Daily index closes: {len(indexes.get('days', [])) if indexes else 0} days")
    nifty_emas = fetch_nifty_emas();           print(f"  Nifty EMAs: {'OK' if nifty_emas else 'unavailable'}")
    macro      = fetch_macro_snapshot();       print(f"  Macro: {len(macro['items']) if macro else 0} indicators")
    sectors    = fetch_sector_summary();       print(f"  Sector perf: top/bottom of {sectors['all_count'] if sectors else 0}")
    flips      = fetch_quadrant_flips();       print(f"  Quadrant flips: {len(flips)}")
    breadth    = fetch_breadth_summary();      print(f"  Breadth: {len(breadth)} days")
    movers     = fetch_top_movers();           print(f"  Top movers: {len(movers['rank_surges']) if movers else 0} surges, {len(movers['vol_breakouts']) if movers else 0} vol")
    earnings   = fetch_earnings_season();      print(f"  Earnings: season={earnings.get('season_period')}, in-season={earnings.get('stocks_in_season')}, stale={earnings.get('stocks_with_stale_data')}")
    signals    = fetch_signals_summary();      print(f"  Signals: {signals.get('total_unique_this_week', 0)} unique stocks across {len(signals.get('by_type', []))} indicators")

    print("\n📰 Fetching Google News headlines...")
    news = fetch_google_news()
    print(f"  Got {len(news)} unique headlines")

    print("\n🧠 Extracting market-moving events (Claude Haiku)...")
    news_digest = extract_news_events(news)
    print(f"  {'Event digest built' if news_digest else 'No digest — falling back to raw headlines'}")

    print("\n🧱 Building context block...")
    context = build_context(macro, indexes, nifty_emas, sectors, flips, breadth, movers, earnings, signals, news, news_digest)
    os.makedirs(OUT_DIR, exist_ok=True)
    today_str = date.today().strftime('%Y-%m-%d')
    ctx_path = f'{OUT_DIR}/_context_{today_str}.md'
    with open(ctx_path, 'w', encoding='utf-8') as f:
        f.write(context)
    print(f"  Context: {len(context):,} chars (~{len(context)//4:,} tokens) -> {ctx_path}")

    print("\n🤖 Calling Claude Opus (long-form)...")
    long_md = call_claude(LONG_FORM_PROMPT, context, max_tokens=6500)
    print(f"  Got {len(long_md):,} chars")

    print("\n🤖 Calling Claude Opus (short-form)...")
    short_md = call_claude(SHORT_FORM_PROMPT, context, max_tokens=2000)
    print(f"  Got {len(short_md):,} chars")

    print("\n💾 Writing reports...")
    paths = {
        'long_md':  f'{OUT_DIR}/{today_str}_market-intel_long.md',
        'short_md': f'{OUT_DIR}/{today_str}_market-intel_short.md',
        'long_dx':  f'{OUT_DIR}/{today_str}_market-intel_long.docx',
        'short_dx': f'{OUT_DIR}/{today_str}_market-intel_short.docx',
    }
    with open(paths['long_md'],  'w', encoding='utf-8') as f: f.write(long_md)
    with open(paths['short_md'], 'w', encoding='utf-8') as f: f.write(short_md)
    md_to_docx(long_md,  paths['long_dx'],  title=f"Weekly Market Intelligence — {today_str}")
    md_to_docx(short_md, paths['short_dx'], title=f"Weekly Market Brief — {today_str}")
    for p in paths.values():
        print(f"  ✅ {p}")

    print(f"\nCompleted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


if __name__ == '__main__':
    main()
