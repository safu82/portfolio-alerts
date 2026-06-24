"""One-shot script to build the Algo (Paper) Trading reference Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\Sarfaraz Khimani\Documents\portfolio-alerts-new\algo-paper-trading-overview.docx"

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Base styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    r = p.add_run(text)
    return p


def kv_bullet(key, value, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    p.add_run(value)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)


def table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row):
            rc[i].text = str(v)
            for p in rc[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if widths_cm:
        for row in t.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    return t


# ────────────────────────────────────────────────────────────────────────
# TITLE
# ────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Algo (Paper) Trading — Full System Reference")
r.bold = True
r.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Strategy v3 — Conviction Stack | D0 scan / D1 fill | Updated 2026-05-14")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()

# ────────────────────────────────────────────────────────────────────────
# 1. OVERVIEW
# ────────────────────────────────────────────────────────────────────────
h1("1. What This Is")
para(
    "The Algo tab on the dashboard runs a fully simulated trading strategy on top of the "
    "existing signal infrastructure. No real broker orders are placed — every entry, partial, "
    "and exit is logged to the paper_trades table in Supabase with mode='paper'. The point is "
    "to measure whether the conviction-based signal stack would have made money under a "
    "disciplined position-sizing and exit framework, with no look-ahead bias."
)
para("Key invariants:", bold=True)
bullet("Sleeve = ₹25,00,000 fixed (no compounding into position sizing).")
bullet("Single user, single instance — runs on GitHub Actions weekday 22:00 IST.")
bullet("D0 scan / D1 fill split: signals are scanned on D0 evening; entry is priced at D1 open via Railway worker.")
bullet("Real holdings are excluded from candidates so the algo never duplicates a position already in the portfolio.")

# ────────────────────────────────────────────────────────────────────────
# 2. DAILY TIMELINE
# ────────────────────────────────────────────────────────────────────────
h1("2. Daily Timeline (IST)")
para(
    "Three distinct processes touch paper_trades each weekday. Below is what happens, when, "
    "where it runs, and what state it leaves behind."
)
table(
    headers=["Time (IST)", "Where", "Job", "What it does"],
    rows=[
        ["09:15 – 15:30", "Railway worker (zerodha_rest_updater_railway.py)",
         "Live price loop (5-sec tick)",
         "Streams LTPs for portfolio + active paper trades into live_prices."],
        ["~09:15:15", "Railway worker, inside live-price loop",
         "fill_pending_paper_trades()",
         "Reads pending rows from D0; uses first successful tick as the D1 'open'; "
         "recomputes entry_price (× 1.0015 slippage), qty, initial_stop with the captured D0 ATR; "
         "flips status='pending' → 'open' with entry_date = today."],
        ["09:15 – 15:30", "Railway worker, every 5-sec tick",
         "check_paper_stops()",
         "For each open paper trade, if live_prices.price ≤ current_stop → close at LTP "
         "(exit_reason='stop' or 'trail_stop'). Mirrors _close_trade field-for-field."],
        ["~15:30", "End of session", "—",
         "live_prices stops updating. From here until the next 09:15, LTP shown on dashboard "
         "falls back to daily_stock_snapshots.close."],
        ["~22:00 (16:30 UTC)", "GitHub Actions: paper-trading.yml", "paper_trader.py main()",
         "1) Process EOD exits on currently open trades using today's OHLC bar (partials, "
         "trailing stops, time stops, universal 25% book). "
         "2) Reload state; compute kill switch. "
         "3) If alive, score today's signals into T1/T2/T3/T4 buckets, apply universal gates "
         "(sector + earnings + cooldown), insert top-ranked candidates as status='pending'. "
         "4) Write paper_equity row + paper_run_log row."],
    ],
    widths_cm=[2.4, 3.2, 4.2, 7.0],
)

para(
    "Auto-cancel of stale pending: the fill job expires any pending row whose scan_date is "
    "more than 2 trading days behind today (exit_reason='fill_expired'). A live tick missing "
    "for a fresh pending row does NOT expire it on day 1 — the row waits for tomorrow.",
    italic=True,
)

# ────────────────────────────────────────────────────────────────────────
# 3. STRATEGY: THE CONVICTION STACK
# ────────────────────────────────────────────────────────────────────────
h1("3. Strategy: The Conviction Stack")
para(
    "The entry universe is ranked into 4 priority buckets. T1 wins ties over T2, T2 wins over "
    "T3, T3 wins over T4. Within a bucket, ties are broken by signal_score (a per-bucket "
    "ranking heuristic — not directly comparable across buckets)."
)

h2("Tier definitions")
table(
    headers=["Tier", "Source", "Trigger", "Risk %", "Position cap"],
    rows=[
        ["T1_MULTI_STRONG", "entry_signals",
         "≥2 distinct STRONG signal_types from {blue_zone_strong, golden_cross_strong, "
         "macd_strong, bull_flag_strong, vcp_strong}. (signal-type level — two of the same "
         "strong type does NOT qualify.)",
         "1.50% of sleeve = ₹37,500", "₹3,00,000"],
        ["T2_STRONG_REG", "entry_signals",
         "≥2 distinct entry_signal FAMILIES, exactly one of which is strong, ≥1 regular. "
         "(Family dedupe: blue_zone_buy + blue_zone_strong counts as 1 BLUE_ZONE family.)",
         "1.00% = ₹25,000", "₹2,50,000"],
        ["T3_MULTI_REG", "entry_signals",
         "≥2 distinct entry_signal FAMILIES, none strong. Pure multi-regular confirmation.",
         "0.66% = ₹16,500", "₹1,50,000"],
        ["T4_RS_ACCEL", "presignal_scores",
         "rule_rs_acceleration == 100 AND any of {rule_bz_buy, rule_bz_strong, "
         "rule_golden_cross, rule_macd, rule_pullback_bounce} > 80 AND rank_slope ≤ –15 AND "
         "rs_30d_high = true, ALL on the same presignal row.",
         "0.50% = ₹12,500", "₹1,00,000"],
    ],
    widths_cm=[3.0, 2.5, 7.0, 2.2, 2.3],
)

h2("Family map (entry_signals)")
code_block(
    "blue_zone_buy        → BLUE_ZONE\n"
    "blue_zone_strong     → BLUE_ZONE\n"
    "golden_cross_buy     → GOLDEN_CROSS\n"
    "golden_cross_strong  → GOLDEN_CROSS\n"
    "macd_buy             → MACD\n"
    "macd_strong          → MACD\n"
    "bull_flag_buy        → BULL_FLAG\n"
    "bull_flag_strong     → BULL_FLAG\n"
    "vcp_buy              → VCP\n"
    "vcp_strong           → VCP\n"
    "narrow_cpr_breakaway → CPR_BREAKAWAY   (regular only, not in strong set)\n"
    "darvas_box           → DARVAS          (regular only, not in strong set)"
)
para(
    "Note: narrow_cpr_breakaway and darvas_box are explicitly NOT treated as STRONG for "
    "Bucket A. They can still contribute to T2/T3 as regular families. This matches the "
    "dashboard's Multi-Strong Buy badge logic exactly.",
    italic=True,
)

h2("signal_score (intra-bucket tiebreak)")
bullet("T1: strong_types_count × 100.")
bullet("T2/T3: families_count × 50 + strong_families × 30.")
bullet("T4: 100 + max(partner-rule score), where partner-rule is the highest-scoring rule above 80.")

# ────────────────────────────────────────────────────────────────────────
# 4. UNIVERSAL GATES
# ────────────────────────────────────────────────────────────────────────
h1("4. Universal Gates (all tiers must pass)")
para("Applied AFTER tier classification, in this order. First failure short-circuits.")

h2("4.1 Hygiene checks")
bullet("Not in real holdings (excluded by ticker against holdings table where portfolio='INDIAN').")
bullet("Not already in an open or pending paper position.")
bullet("Cooldown: not closed within the last 21 trading days (window loaded for 28 calendar days for safety).")

h2("4.2 Sector quadrant (RRG)")
para(
    "The stock's sector (from indian_stock_sectors) must be in the LEADING or IMPROVING "
    "quadrant of the Relative Rotation Graph. Sector quadrants come from the "
    "get_all_sector_rankings RPC."
)
bullet("Allowed: leading, improving.")
bullet("Rejected: weakening, lagging, sector_unknown, sector_quadrant_missing.")

h2("4.3 Earnings momentum")
para(
    "Looks at the quarterly_financials JSON column on stock_fundamentals (newest first). "
    "Requires ≥6 quarters because the test compares latest YoY against the prior-quarter YoY, "
    "and each YoY needs the same quarter from a year ago."
)
para("Math:", bold=True)
code_block(
    "latest_yoy_rev  = (Q[0].revenue_cr   - Q[4].revenue_cr)   / Q[4].revenue_cr\n"
    "prior_yoy_rev   = (Q[1].revenue_cr   - Q[5].revenue_cr)   / Q[5].revenue_cr\n"
    "latest_yoy_np   = (Q[0].net_income_cr - Q[4].net_income_cr) / Q[4].net_income_cr\n"
    "prior_yoy_np    = (Q[1].net_income_cr - Q[5].net_income_cr) / Q[5].net_income_cr\n\n"
    "PASS if:\n"
    "  latest_yoy_rev  > 0   AND  latest_yoy_rev  > prior_yoy_rev\n"
    "  latest_yoy_np   > 0   AND  latest_yoy_np   > prior_yoy_np"
)
para(
    "Reject reasons surfaced in the funnel: quarters_<N>, revenue_yoy_not_positive, "
    "revenue_yoy_not_improving, profit_yoy_not_positive, profit_yoy_not_improving, "
    "revenue_data_incomplete, profit_data_incomplete, fundamentals_missing.",
    italic=True,
)

h2("4.4 Portfolio limits")
bullet("Max 8 open positions at once (open + pending count toward this cap).")
bullet("Max 3 new entries per day (after kill-switch + capacity checks).")
bullet("Sector concentration: notional in any one sector ≤ 25% of sleeve = ₹6,25,000.")
bullet("Position floor: notional ≥ ₹40,000. If sizing produces a smaller notional, the trade is rejected.")

h2("4.5 Kill switches")
para("Computed before opening new entries. Either trigger pauses new entries for that day; exits and trail updates still run.")
bullet("Daily kill: provisional total < previous total_value by more than 2%.")
bullet("Drawdown kill: provisional total below all-time-high by more than 8%.")

# ────────────────────────────────────────────────────────────────────────
# 5. POSITION SIZING
# ────────────────────────────────────────────────────────────────────────
h1("5. Position Sizing")
para("Two-constraint sizing: dollar risk AND notional cap. Whichever produces the smaller qty wins.")
code_block(
    "risk_inr    = SLEEVE × tier.risk_pct\n"
    "stop_dist   = STOP_ATR_MULT × ATR_14   # STOP_ATR_MULT = 2.0\n"
    "qty_by_risk = floor(risk_inr / stop_dist)\n"
    "qty_by_cap  = floor(tier.cap  / entry_price)\n"
    "qty         = max(0, min(qty_by_risk, qty_by_cap))\n\n"
    "if qty × entry_price < ₹40,000:  REJECT (position_floor)"
)
para("Per-tier numbers at ₹25L sleeve:")
table(
    headers=["Tier", "Dollar risk", "Notional cap", "Effective max @ entry"],
    rows=[
        ["T1_MULTI_STRONG", "₹37,500", "₹3,00,000", "Limited by notional cap on most names."],
        ["T2_STRONG_REG",   "₹25,000", "₹2,50,000", "Cap binds on lower-vol names; risk binds on high-ATR names."],
        ["T3_MULTI_REG",    "₹16,500", "₹1,50,000", "Risk usually binds (smaller dollar risk)."],
        ["T4_RS_ACCEL",     "₹12,500", "₹1,00,000", "Tight cap; meant to be the smallest probe-sized bucket."],
    ],
    widths_cm=[3.5, 2.5, 2.5, 7.0],
)

para("Slippage: entry_price = first D1 live tick × 1.0015 (15 bps).", italic=True)

# ────────────────────────────────────────────────────────────────────────
# 6. EXIT LOGIC
# ────────────────────────────────────────────────────────────────────────
h1("6. Exit Logic")
para(
    "Exits run in two places: intraday (Railway, on every 5-sec live tick) handles only the "
    "hard stop; everything else — partials, trailing stop updates, universal 25%/15-day book, "
    "time stop — runs in the EOD 22:00 IST batch using today's daily_stock_snapshots bar."
)

h2("6.1 Initial stop")
para("Set at fill time: initial_stop = entry_price − 2 × ATR_14 (using D0 ATR captured in entry_atr).")

h2("6.2 Partial booking (R-multiples)")
para("R = entry_price − initial_stop = per-share risk. Partial qty = 33% of initial_quantity (rounded down, min 1).")
table(
    headers=["Tier", "1st partial at", "2nd partial at", "After 1st partial", "After 2nd partial"],
    rows=[
        ["T1", "+3R", "+6R", "Stop → breakeven (entry_price)", "trail_armed = true; trailing 2×ATR from close kicks in"],
        ["T2", "+2R", "+4R", "Stop → breakeven", "trail_armed = true"],
        ["T3", "+2R", "+4R", "Stop → breakeven", "trail_armed = true"],
        ["T4", "+2R", "+4R", "Stop → breakeven", "trail_armed = true"],
    ],
    widths_cm=[1.5, 2.5, 2.5, 4.5, 5.0],
)
para(
    "When the second partial fires, the residual ~33% rides with a 2×ATR trailing stop "
    "(updated only when new_stop > current_stop — never ratchets down). If both partial "
    "targets are taken inside the same bar, both are recorded with their respective target "
    "prices in the partials JSON.",
    italic=True,
)

h2("6.3 Universal 25% / 15-day book")
para(
    "Independent of R-partials, to avoid letting a fast vertical move round-trip. Triggered "
    "only when partials_taken == 0 AND holding_days < 15 trading days AND bar_close has gained "
    "≥25% from entry. Books 25% of initial_quantity at bar_close and arms breakeven."
)

h2("6.4 Trailing stop")
para(
    "Armed when partials_taken reaches 2, OR early via the +10% MFE safety net "
    "(EARLY_TRAIL_PCT) — whichever comes first sets trail_armed = true and locks the stop "
    "at breakeven. Each EOD thereafter the stop is raised to the HIGHER of: (a) the 2×ATR "
    "trail, bar_close − 2 × today's ATR_14; and (b) a give-back floor at entry_price + 0.5 × "
    "(peak open gain), anchored to the MFE high-water mark (GIVEBACK_LOCK_FRAC = 0.5). The "
    "floor caps round-trip on a runner to half its peak gain and protects fast reversals that "
    "the close-following ATR trail misses. The stop only ratchets up, never down; the give-back "
    "floor books no profit, so r_multiple stays clean."
)

h2("6.5 Time stop")
para(
    "After 25 trading days, if the trade has NOT armed its trailing stop — i.e. it reached "
    "neither its 2nd R-partial nor a +10% MFE early-trail — close the residual qty at bar_close "
    "with exit_reason='time_stop'. The cut is P&L-agnostic: a position drifting at −4% or +6% "
    "that never reached a trail-arming target is dead money and gets recycled. The check runs in "
    "process_exits AFTER the partial/trail/universal-book pass, so any position that has armed its "
    "trail (and is therefore genuinely running) is exempt."
)

h2("6.6 Intraday stop (Railway)")
para(
    "Runs every 5 seconds during market hours on the live tick: if live_prices.price ≤ "
    "current_stop for any open paper trade, close the remaining qty at LTP. Sets "
    "exit_reason='trail_stop' if trail_armed, else 'stop'. Mirrors _close_trade so the EOD "
    "22:00 batch sees a clean closed row."
)

h2("6.7 Exit reasons cheat-sheet")
table(
    headers=["exit_reason", "Where", "What it means"],
    rows=[
        ["stop", "Railway intraday OR EOD batch", "Initial stop hit (before any partial or trail)."],
        ["trail_stop", "Railway intraday OR EOD batch", "Stop hit AFTER trail was armed (i.e. after 2nd partial)."],
        ["partials_full", "EOD batch only", "Both partial targets filled the same bar; no residual."],
        ["time_stop", "EOD batch", "25 trading days held with trailing stop never armed (no 2nd partial / +10% MFE)."],
        ["fill_expired", "Railway fill job", "Pending row aged > 2 trading days with no live tick."],
        ["fill_rejected_missing_atr", "Railway fill job", "entry_atr was null/0 on the pending row (shouldn't happen)."],
        ["fill_rejected_position_floor", "Railway fill job",
         "D1-open price moved enough that sizing would produce qty=0 or notional < ₹40k."],
    ],
    widths_cm=[3.8, 4.0, 8.0],
)

# ────────────────────────────────────────────────────────────────────────
# 7. DATA MODEL
# ────────────────────────────────────────────────────────────────────────
h1("7. Data Model (Supabase)")
para("All paper-trade state lives in three tables, all keyed on mode='paper'.")

h2("paper_trades")
para("One row per simulated trade. Lifecycle: pending → open → closed.")
code_block(
    "id                    bigserial PK\n"
    "ticker                text\n"
    "sector                text\n"
    "strategy_tier         CHECK in {'T1_MULTI_STRONG','T2_STRONG_REG','T3_MULTI_REG','T4_RS_ACCEL'}\n"
    "signal_families       text[]      -- e.g. ['BLUE_ZONE','MACD']\n"
    "strong_family_count   int\n"
    "signal_ids            jsonb       -- {'presignal_ids':[...], 'entry_signal_ids':[...]}\n"
    "signal_score          numeric\n"
    "entry_date            date        -- D0 scan date until filled; D1 fill date after\n"
    "entry_price           numeric     -- D0 placeholder until filled; real D1 entry after\n"
    "entry_atr             numeric     -- D0 ATR_14, captured at scan time, never overwritten\n"
    "initial_quantity      int\n"
    "current_quantity      int         -- decremented by partials and stops\n"
    "entry_value           numeric     -- qty × entry_price\n"
    "initial_risk          numeric     -- qty × (entry − stop)\n"
    "initial_stop          numeric\n"
    "current_stop          numeric     -- ratcheted up by breakeven + trail\n"
    "breakeven_armed       bool        -- true after 1st partial\n"
    "trail_armed           bool        -- true after 2nd partial\n"
    "partials_taken        int         -- 0, 1, or 2\n"
    "partials              jsonb       -- list of {date, kind, price, qty, pnl}\n"
    "realised_pnl          numeric     -- running, includes partials\n"
    "realised_qty          int         -- running\n"
    "exit_date             date        -- null until closed\n"
    "exit_price            numeric\n"
    "exit_reason           text\n"
    "total_pnl             numeric\n"
    "total_pnl_pct         numeric\n"
    "r_multiple            numeric     -- realised_pnl_per_initial_share / risk_per_share\n"
    "holding_days          int         -- trading days, not calendar days\n"
    "status                CHECK in {'pending','open','closed'}\n"
    "mode                  text        -- 'paper'\n"
    "created_at            timestamptz\n"
    "updated_at            timestamptz"
)

h2("paper_equity")
para("One row per scan day. Upserted by paper_trader.write_equity_row.")
code_block(
    "snapshot_date         date PK\n"
    "cash                  numeric    -- SLEEVE − sum(entry_value of open) + cum_realised_closed\n"
    "positions_value       numeric    -- sum(close × current_qty) across open\n"
    "total_value           numeric    -- cash + positions_value\n"
    "drawdown_pct          numeric    -- (total_value − ATH) / ATH × 100\n"
    "open_positions        int\n"
    "trades_opened_today   int\n"
    "trades_closed_today   int\n"
    "realised_today        numeric"
)

h2("paper_run_log")
para("One row per scheduled run, including failed ones. Powers the Latest Run card on the dashboard.")
code_block(
    "id                    bigserial PK\n"
    "run_date              date\n"
    "signals_seen          int        -- presignal + entry_signal rows in window\n"
    "signals_passed        int        -- qualified after all gates\n"
    "signals_taken         int        -- actually inserted as pending\n"
    "exits_processed       int\n"
    "kill_switch_active    bool\n"
    "kill_switch_reason    text\n"
    "funnel                jsonb      -- bucket counts + per-reason rejection counts\n"
    "errors                jsonb      -- null on a clean run"
)

# ────────────────────────────────────────────────────────────────────────
# 8. DASHBOARD (Algo tab)
# ────────────────────────────────────────────────────────────────────────
h1("8. Dashboard — Algo Tab")
para("Lives in index.html, rendered by renderPaperTradingTab(). Active when activeTab === 'paper'.")

h2("Layout, top to bottom")
bullet("Tier filter strip: ALL | T1 Multi-Strong | T2 Strong+Reg | T3 Multi-Reg | T4 RS Accel.")
bullet("KPI strip (6 cards): Sleeve Value, Open/Closed, Win Rate, Avg R, Realised P&L, Max Drawdown.")
bullet("Equity curve: paper_equity total_value vs ₹25L baseline.")
bullet("Pending Fills card: rows where status='pending' from D0 scan, awaiting D1 fill.")
bullet("Open Positions table: ticker, families, entry, qty, entry ₹, LTP, MTM, unr %, current stop (% from entry), partials/breakeven/trail flags, realised, days held.")
bullet("Closed Trades table: entry→exit, prices, P&L, P&L %, R, exit_reason, days held.")
bullet("Latest Run card: funnel from paper_run_log (signals seen → qualified → taken, plus exits and rejection reasons).")

h2("LTP source for the dashboard")
para(
    "Priority: (1) live_prices via the global prices object (populated during market hours by "
    "the Railway worker for both real holdings and active paper trades); (2) fallback "
    "paperLatestClose loaded from daily_stock_snapshots.close on the most recent equity "
    "snapshot date. Outside market hours and weekends, you'll see EOD closes."
)

# ────────────────────────────────────────────────────────────────────────
# 9. HELPER SCRIPTS
# ────────────────────────────────────────────────────────────────────────
h1("9. Helper Scripts")
table(
    headers=["Script", "What it does", "Writes?"],
    rows=[
        ["paper_trader.py", "Main EOD executor. Runs in GitHub Actions at 22:00 IST.", "Yes — paper_trades, paper_equity, paper_run_log."],
        ["paper_trader_dryrun.py", "Same entry pipeline read-only. Prints funnel, ranked candidates, sample rejections.", "No."],
        ["paper_fill_pending.py", "Manual CLI safety net for the D1 fill (kept in case Railway is down).", "Yes — flips pending → open."],
        ["list_open_trades.py", "Prints currently open paper trades with tier, sector, qty, stop, families.", "No."],
        ["check_paper_footprint.py", "Row counts across paper_trades / paper_equity / paper_run_log + lists closed trades.", "No."],
        ["wipe_paper_data.py", "Destructive: clears all three paper tables. Used 2026-05-13 to reset post-v3 cutover.", "Yes — DELETE."],
    ],
    widths_cm=[4.0, 8.5, 4.5],
)

# ────────────────────────────────────────────────────────────────────────
# 10. INPUTS / DEPENDENCIES
# ────────────────────────────────────────────────────────────────────────
h1("10. Inputs (read by paper_trader.py)")
bullet("daily_stock_snapshots — today's OHLC + ATR_14 for both exits and entry sizing.")
bullet("entry_signals — for T1/T2/T3 classification (filtered by alert_date = today).")
bullet("presignal_scores — for T4 classification (filtered by score_date = today).")
bullet("holdings — to exclude tickers already in the real INDIAN portfolio.")
bullet("indian_stock_sectors — ticker → sector.")
bullet("stock_fundamentals.quarterly_financials — earnings YoY filter.")
bullet("get_all_sector_rankings RPC — sector RRG quadrants. Returns {as_of, total_sectors, sectors:{name: {rrg_quadrant,...}}} where sectors is a DICT keyed by sector name.")
bullet("paper_trades — read for open/pending/closed-cooldown state.")
bullet("paper_equity — read recent rows for kill-switch baseline + ATH calc.")

# ────────────────────────────────────────────────────────────────────────
# 11. KNOWN GOTCHAS
# ────────────────────────────────────────────────────────────────────────
h1("11. Known Gotchas (operator notes)")
bullet(
    "EARNINGS_MIN_QUARTERS = 6, not 5. The 'YoY improving' test compares latest YoY against "
    "prior-quarter YoY, so it needs Q[0]..Q[5]. Don't drop to 5 without redoing the math."
)
bullet(
    "paper_trades.strategy_tier has a CHECK constraint listing all 4 tiers. Adding a new tier "
    "requires an ALTER on Supabase first, or inserts crash with code 23514."
)
bullet(
    "Family dedupe varies by bucket. T1 counts distinct strong signal_TYPES (so two "
    "blue_zone_strong rows don't qualify). T2/T3 dedupe to FAMILY (blue_zone_buy + "
    "blue_zone_strong = 1 BLUE_ZONE family). T4 is presignal-only with hardcoded gates."
)
bullet(
    "narrow_cpr_breakaway and darvas_box are NOT in the STRONG set for T1, even though they "
    "appear in entry_signals. They can still contribute as regular families in T2/T3."
)
bullet(
    "get_recent_trading_dates uses a single liquid benchmark (RELIANCE.NS first, then HDFCBANK "
    "etc.) instead of paginating the full snapshots table — PostgREST's 1000-row response cap "
    "silently truncated the signal lookback window in the old approach."
)
bullet(
    "GitHub Actions cron drifts 10–30 min. That's why the fill moved to the Railway worker, "
    "which fires the moment the first 09:15 tick lands. The previous GH-Actions fill missed "
    "the actual open."
)
bullet(
    "paper_trades.entry_date means 'scan date' while status='pending' and 'fill date' once "
    "status='open'. Same column, different meaning by status. Cooldown logic uses exit_date."
)
bullet(
    "The signal window in paper_trader.py is D0-only (today's signals only). Earlier versions "
    "used a 1-2 day window which produced look-ahead bias when combined with D0-close entry "
    "pricing."
)

# ────────────────────────────────────────────────────────────────────────
# 12. CHANGE LOG (v3 cutover)
# ────────────────────────────────────────────────────────────────────────
h1("12. Recent Change Log")
table(
    headers=["Date", "Change", "Commit"],
    rows=[
        ["2026-05-13", "D0-only signal window — remove the residual 1-day lookback that allowed look-ahead bias.", "fdbd934"],
        ["2026-05-13", "Move paper-trade fill from GitHub Actions to the Railway worker — fires inside the live-price loop at the first 09:15 tick.", "5fbd60b"],
        ["2026-05-13", "Add D1-open fills + intraday hard-stop monitor on Railway.", "d157156"],
        ["2026-05-13", "Paper trader v3 — 4-bucket conviction stack (T1 > T2 > T3 > T4) + universal sector & earnings gates.", "51de0f1"],
        ["2026-05-13", "Fix signal-window pagination bug (PostgREST 1000-row cap silently truncated lookback).", "5ba9a68"],
        ["2026-05-14", "First verified D1 fills via Railway at 09:15. System working end-to-end.", "(operational)"],
    ],
    widths_cm=[2.5, 11.5, 3.0],
)

# ────────────────────────────────────────────────────────────────────────
# 13. QUICK REFERENCE
# ────────────────────────────────────────────────────────────────────────
h1("13. Quick Reference")
para("All tunable constants from paper_trader.py:")
code_block(
    "SLEEVE                = 2_500_000   # ₹25L paper sleeve\n"
    "COOLDOWN_DAYS         = 21          # trading days after a close\n"
    "POSITION_FLOOR        = 40_000      # min notional, ₹\n"
    "SLIPPAGE_BPS          = 15          # 0.15% on entry\n"
    "STOP_ATR_MULT         = 2.0         # initial stop = 2 × ATR_14\n"
    "TRAIL_ATR_MULT        = 2.0         # trailing stop = 2 × ATR\n"
    "EARLY_TRAIL_PCT       = 10.0        # arm trail early if MFE hits +10%\n"
    "GIVEBACK_LOCK_FRAC    = 0.5         # floor stop at entry + 50% of peak MFE\n"
    "TIME_STOP_DAYS        = 25          # close at N days if trail never armed\n"
    "UNIVERSAL_BOOK_DAYS   = 15          # window for the 25% fast-move book\n"
    "UNIVERSAL_BOOK_PCT    = 25.0        # trigger %\n"
    "UNIVERSAL_BOOK_QTY_PCT= 25          # book qty %\n"
    "PARTIAL_QTY_PCT       = 33          # R-partial qty %\n"
    "MAX_OPEN_POSITIONS    = 8           # open + pending\n"
    "MAX_SECTOR_CONC       = 0.25        # 25% of sleeve per sector\n"
    "MAX_NEW_PER_DAY       = 3\n"
    "DAILY_KILL_PCT        = -2.0        # pause new entries if day < -2%\n"
    "DD_KILL_PCT           = -8.0        # pause new entries if DD < -8%\n"
    "EARNINGS_MIN_QUARTERS = 6\n"
    "PENDING_MAX_TRADING_DAYS = 2        # Railway expiry threshold\n"
    "RS_ACCEL_REQUIRED_SCORE = 100       # T4 main rule\n"
    "RS_ACCEL_OTHER_RULE_MIN =  80       # T4 partner rule\n"
    "RS_ACCEL_RANK_SLOPE_MAX = -15       # T4 rank slope (negative = improving rank)"
)

# Save
doc.save(OUT)
print(f"Wrote: {OUT}")
