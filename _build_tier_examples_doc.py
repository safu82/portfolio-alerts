"""Algo Paper Trading — worked-numbers tier reference."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\Sarfaraz Khimani\Documents\portfolio-alerts-new\algo-tier-worked-examples.docx"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.add_run(text)


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)


def callout(text, color=(0xee, 0xf2, 0xff)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)


def table(headers, rows, widths_cm=None, header_color=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row):
            rc[i].text = str(v)
            for p in rc[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if widths_cm:
        for row in t.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)


def fmt_inr(n):
    """Format integer rupee with Indian-grouping thousands."""
    s = f"{int(round(n)):,}"
    return f"₹{s}"


# ──────────────────────────────────────────────────────────────────────
# TITLE
# ──────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Algo Paper Trading — Tier-by-Tier Worked Examples")
r.bold = True; r.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Every number, every milestone, every tier — with actual rupees")
r.italic = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────
# MENTAL MODEL
# ──────────────────────────────────────────────────────────────────────
h1("1. The Mental Model")
para("Three numbers drive everything that follows. Memorize these:", bold=True)
code_block(
    "R          = entry_price − initial_stop      (per-share risk in ₹)\n"
    "stop_dist  = 2 × ATR_14                       (= 1R, by construction)\n"
    "SLEEVE     = ₹25,00,000                       (fixed paper capital)"
)
para(
    "Because stop_dist is fixed at 2×ATR, R is always equal to 2×ATR. So every "
    "R-multiple in the strategy translates directly into an ATR multiple of distance "
    "from the entry price:"
)
table(
    headers=["R-multiple", "In ATR units", "Meaning"],
    rows=[
        ["1R", "2 × ATR", "Per-share risk. Distance from entry to initial stop."],
        ["2R", "4 × ATR", "First partial target for T2/T3/T4."],
        ["3R", "6 × ATR", "First partial target for T1."],
        ["4R", "8 × ATR", "Second partial target for T2/T3/T4."],
        ["6R", "12 × ATR", "Second partial target for T1."],
    ],
    widths_cm=[2.5, 3.0, 11.5],
)

para(
    "Position size is constrained by two limits — whichever produces the SMALLER qty wins. "
    "For each tier we keep both numbers handy:"
)
code_block(
    "qty_by_risk = floor( SLEEVE × risk_pct / stop_dist )    # dollar risk constraint\n"
    "qty_by_cap  = floor( cap / entry_price )                # notional cap constraint\n"
    "qty         = min(qty_by_risk, qty_by_cap)\n"
    "REJECT if qty × entry_price < ₹40,000  (position floor)"
)
para("Per-tier risk_pct and cap:")
table(
    headers=["Tier", "risk_pct", "Dollar risk @ ₹25L", "Notional cap"],
    rows=[
        ["T1_MULTI_STRONG", "1.50%", "₹37,500", "₹3,00,000"],
        ["T2_STRONG_REG",   "1.00%", "₹25,000", "₹2,50,000"],
        ["T3_MULTI_REG",    "0.66%", "₹16,500", "₹1,50,000"],
        ["T4_RS_ACCEL",     "0.50%", "₹12,500", "₹1,00,000"],
    ],
    widths_cm=[4.5, 2.5, 4.0, 4.0],
)

# ──────────────────────────────────────────────────────────────────────
# SHARED SETUP
# ──────────────────────────────────────────────────────────────────────
h1("2. Shared Worked-Example Setup")
para(
    "To make tiers directly comparable, every worked example below uses the SAME "
    "underlying stock. The only thing that varies between tiers is risk_pct, cap, "
    "and the R-multiple ladder."
)
code_block(
    "Hypothetical stock: XYZ.NS\n"
    "  D1 09:15 first tick     = ₹499.25\n"
    "  Slippage 15bps          → entry_price = ₹500.00\n"
    "  ATR_14 (from D0)        = ₹15.00       (3% of price — a typical mid/large cap)\n"
    "  stop_dist  (2 × ATR)    = ₹30.00\n"
    "  initial_stop            = ₹470.00\n"
    "  R                       = ₹30 per share"
)
para("Five exit scenarios are tested for every tier so you can see the full distribution:", bold=True)
bullet("Scenario A — Trend works: trade hits +3R / +6R (T1) or +2R / +4R (others) and the residual trails up to +9R before stopping out.")
bullet("Scenario B — Pop then reverse: 1st partial fills, breakeven stop catches the pullback before 2nd partial.")
bullet("Scenario C — Stopped cold: price drops to initial stop before any partial. The worst case.")
bullet("Scenario D — Vertical pop: +25% in the first 15 trading days, no R-partial yet → universal 25% book triggers.")
bullet("Scenario E — Time stop: 25 trading days of dead money (price drifts in [−2%, +2%]).")


def render_tier(
    tier_name,
    tier_label,
    risk_pct,
    risk_pct_pretty,
    cap,
    partial_R,           # (R1, R2)
    sleeve=2_500_000,
    entry=500.0,
    atr=15.0,
    stop_dist=30.0,
    stop=470.0,
    R=30.0,
):
    """Render one full tier section with numbers."""
    h1(f"3. {tier_name} — {tier_label}")

    # ─── Sizing ───
    h2("Position sizing")
    risk_inr = sleeve * risk_pct
    qty_by_risk = int(risk_inr // stop_dist)
    qty_by_cap = int(cap // entry)
    qty = max(0, min(qty_by_risk, qty_by_cap))
    bind = "notional cap" if qty == qty_by_cap and qty < qty_by_risk else (
        "dollar risk" if qty == qty_by_risk and qty < qty_by_cap else "tie")
    notional = qty * entry
    actual_risk = qty * R

    code_block(
        f"risk_inr     = ₹25,00,000 × {risk_pct_pretty}  = {fmt_inr(risk_inr)}\n"
        f"qty_by_risk  = floor( {fmt_inr(risk_inr)} / ₹30 )  = {qty_by_risk}\n"
        f"qty_by_cap   = floor( {fmt_inr(cap)} / ₹500 )       = {qty_by_cap}\n"
        f"qty          = min({qty_by_risk}, {qty_by_cap})  = {qty}   ← {bind} binds\n"
        f"notional     = {qty} × ₹500  = {fmt_inr(notional)}\n"
        f"actual_risk  = {qty} × ₹30   = {fmt_inr(actual_risk)}   ({actual_risk/sleeve*100:.2f}% of sleeve)"
    )
    if bind == "notional cap":
        para(
            f"Because the cap binds, the trade actually risks {fmt_inr(actual_risk)} "
            f"({actual_risk/sleeve*100:.2f}%), which is LESS than the {risk_pct_pretty} target. "
            "This is by design — the cap is a hard ceiling on tier exposure for low-ATR names.",
            italic=True,
        )

    # ─── Stop & R ladder ───
    h2("Stop loss & R-multiple ladder")
    R1, R2 = partial_R
    t1_price = entry + R1 * R
    t2_price = entry + R2 * R
    partial_qty = max(1, int(qty * 33 / 100))
    residual_qty = qty - 2 * partial_qty

    code_block(
        f"initial_stop  = ₹500 − (2 × ₹15)  = ₹470.00\n"
        f"R             = ₹500 − ₹470       = ₹30 per share\n\n"
        f"1st partial:  +{R1}R  → ₹500 + ({R1} × ₹30) = ₹{t1_price:.2f}\n"
        f"2nd partial:  +{R2}R  → ₹500 + ({R2} × ₹30) = ₹{t2_price:.2f}\n\n"
        f"partial_qty   = floor({qty} × 33%) = {partial_qty} shares (each partial)\n"
        f"residual_qty  = {qty} − 2 × {partial_qty} = {residual_qty} shares (trails)"
    )

    # ─── Scenario A: Trend works ───
    h2("Scenario A — Trend works (best case)")
    para(
        f"Price rallies. At +{R1}R (₹{t1_price:.2f}) the 1st partial fires and the stop "
        f"jumps to ₹500 (breakeven). At +{R2}R (₹{t2_price:.2f}) the 2nd partial fires and "
        f"trail_armed flips on. The residual {residual_qty} shares trail at 2×ATR below "
        "the close. We model the trail catching out at +9R (₹770) — i.e., price ran "
        "to roughly +10R then pulled back one ATR-pair."
    )
    pnl_p1 = (t1_price - entry) * partial_qty
    pnl_p2 = (t2_price - entry) * partial_qty
    trail_exit = entry + 9 * R
    pnl_res = (trail_exit - entry) * residual_qty
    total_A = pnl_p1 + pnl_p2 + pnl_res
    table(
        headers=["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        rows=[
            [f"1st partial fills at +{R1}R", f"₹{t1_price:.2f}", str(partial_qty), fmt_inr(pnl_p1), fmt_inr(pnl_p1)],
            [f"2nd partial fills at +{R2}R", f"₹{t2_price:.2f}", str(partial_qty), fmt_inr(pnl_p2), fmt_inr(pnl_p1 + pnl_p2)],
            ["Trail stop catches residual @ +9R", f"₹{trail_exit:.2f}", str(residual_qty), fmt_inr(pnl_res), fmt_inr(total_A)],
        ],
        widths_cm=[6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario A total P&L: {fmt_inr(total_A)}  (R-multiple ≈ {total_A/actual_risk:.2f}R)")

    # ─── Scenario B: Pop then reverse ───
    h2("Scenario B — Pop then reverse")
    para(
        f"Price hits +{R1}R, the 1st partial fires, stop moves to breakeven. Then the "
        "stock pulls back and trips the breakeven stop before reaching the 2nd partial. "
        f"Residual qty closes at ₹500 (no loss on residual, since breakeven stop = entry)."
    )
    pnl_B_p1 = (t1_price - entry) * partial_qty
    pnl_B_rest = 0  # breakeven on remaining
    total_B = pnl_B_p1 + pnl_B_rest
    table(
        headers=["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        rows=[
            [f"1st partial fills at +{R1}R", f"₹{t1_price:.2f}", str(partial_qty), fmt_inr(pnl_B_p1), fmt_inr(pnl_B_p1)],
            ["Breakeven stop hit", "₹500.00", str(qty - partial_qty), fmt_inr(pnl_B_rest), fmt_inr(total_B)],
        ],
        widths_cm=[6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario B total P&L: {fmt_inr(total_B)}  (R-multiple ≈ {total_B/actual_risk:.2f}R)")

    # ─── Scenario C: Stopped cold ───
    h2("Scenario C — Stopped cold (worst case)")
    para(
        f"Price drops straight to the initial stop. All {qty} shares exit at ₹470. "
        "This is exactly the dollar risk computed at sizing time."
    )
    pnl_C = -actual_risk
    table(
        headers=["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        rows=[
            ["Initial stop hit", "₹470.00", str(qty), fmt_inr(pnl_C), fmt_inr(pnl_C)],
        ],
        widths_cm=[6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario C total P&L: {fmt_inr(pnl_C)}  (R-multiple = −1.00R, by construction)")

    # ─── Scenario D: Vertical pop ───
    h2("Scenario D — Vertical pop (universal 25% rule)")
    para(
        "Price gains +25% (to ₹625) within the first 15 trading days, before any R-target "
        "is hit. Universal rule books 25% of initial_quantity at the day's close and arms "
        "breakeven. Then assume price keeps running and eventually hits the +R2 partial "
        f"target (₹{t2_price:.2f}), then trails out at +9R like Scenario A."
    )
    uni_qty = max(1, int(qty * 25 / 100))
    pnl_uni = (625 - entry) * uni_qty
    # After universal book: remaining qty handles partials. Note: this is a stylized
    # example since the universal book skips the 1st partial; we model 2nd partial
    # firing on the remainder.
    remain_after_uni = qty - uni_qty
    # Per code: partials_taken tracking continues — 1st partial at +R1 not skipped, just
    # the universal book pulled forward. For simplicity model 1st partial on remain_after_uni.
    p1_qty_D = min(partial_qty, remain_after_uni)
    pnl_D_p1 = (t1_price - entry) * p1_qty_D
    remain_after_p1 = remain_after_uni - p1_qty_D
    p2_qty_D = min(partial_qty, remain_after_p1)
    pnl_D_p2 = (t2_price - entry) * p2_qty_D
    res_D = remain_after_p1 - p2_qty_D
    pnl_D_res = (trail_exit - entry) * res_D
    total_D = pnl_uni + pnl_D_p1 + pnl_D_p2 + pnl_D_res
    table(
        headers=["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        rows=[
            ["Universal 25% book", "₹625.00", str(uni_qty), fmt_inr(pnl_uni), fmt_inr(pnl_uni)],
            [f"1st partial at +{R1}R", f"₹{t1_price:.2f}", str(p1_qty_D), fmt_inr(pnl_D_p1), fmt_inr(pnl_uni + pnl_D_p1)],
            [f"2nd partial at +{R2}R", f"₹{t2_price:.2f}", str(p2_qty_D), fmt_inr(pnl_D_p2), fmt_inr(pnl_uni + pnl_D_p1 + pnl_D_p2)],
            ["Trail residual @ +9R", f"₹{trail_exit:.2f}", str(res_D), fmt_inr(pnl_D_res), fmt_inr(total_D)],
        ],
        widths_cm=[6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario D total P&L: {fmt_inr(total_D)}  (R-multiple ≈ {total_D/actual_risk:.2f}R)")

    # ─── Scenario E: Time stop ───
    h2("Scenario E — Time stop (dead money)")
    para(
        "25 trading days pass. Price is at ₹505 — within the [−2%, +2%] = [₹490, ₹510] "
        "band. Time stop fires and closes all remaining qty at close. No partial has "
        "fired (otherwise breakeven_armed would have caught a deeper pullback first), "
        "so the full original qty exits at ₹505."
    )
    pnl_E = (505 - entry) * qty
    table(
        headers=["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        rows=[
            ["Time stop @ day 25", "₹505.00", str(qty), fmt_inr(pnl_E), fmt_inr(pnl_E)],
        ],
        widths_cm=[6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario E total P&L: {fmt_inr(pnl_E)}  (R-multiple ≈ {pnl_E/actual_risk:+.2f}R)")

    # ─── Tier summary ───
    h2("Scenario summary for this tier")
    table(
        headers=["Scenario", "Total P&L", "R-multiple", "% of sleeve"],
        rows=[
            ["A — Trend works", fmt_inr(total_A), f"+{total_A/actual_risk:.2f}R", f"{total_A/sleeve*100:+.2f}%"],
            ["B — Pop then reverse", fmt_inr(total_B), f"+{total_B/actual_risk:.2f}R", f"{total_B/sleeve*100:+.2f}%"],
            ["C — Stopped cold", fmt_inr(pnl_C), "−1.00R", f"{pnl_C/sleeve*100:+.2f}%"],
            ["D — Vertical pop", fmt_inr(total_D), f"+{total_D/actual_risk:.2f}R", f"{total_D/sleeve*100:+.2f}%"],
            ["E — Time stop", fmt_inr(pnl_E), f"{pnl_E/actual_risk:+.2f}R", f"{pnl_E/sleeve*100:+.2f}%"],
        ],
        widths_cm=[6.0, 3.5, 3.0, 3.0],
    )

    return {
        'tier': tier_name, 'qty': qty, 'notional': notional, 'risk': actual_risk,
        'A': total_A, 'B': total_B, 'C': pnl_C, 'D': total_D, 'E': pnl_E,
        't1_price': t1_price, 't2_price': t2_price,
        'partial_qty': partial_qty, 'residual_qty': residual_qty,
    }


# Generate per-tier sections
t1 = render_tier(
    'Tier T1', 'Multi-Strong (highest conviction)',
    risk_pct=0.015, risk_pct_pretty='1.50%',
    cap=300_000, partial_R=(3, 6),
)
# Renumber subsequent H1s by hand: we use 4, 5, 6 for T2/T3/T4
# but the render_tier sets H1 to "3. ..."; let's tweak by setting a flag.


def render_tier_n(n, *args, **kwargs):
    # Wrapper to override the H1 number
    # Simplest path: call render_tier and the doc already has the section.
    pass


# Easier: rewrite sections T2, T3, T4 inline using the same helper but with manual H1 numbers.
def render_tier_numbered(section_no, tier_name, tier_label, risk_pct, risk_pct_pretty, cap, partial_R):
    """Same as render_tier but with explicit section number in the heading."""
    sleeve = 2_500_000; entry = 500.0; atr = 15.0
    stop_dist = 30.0; stop = 470.0; R = 30.0
    h1(f"{section_no}. {tier_name} — {tier_label}")

    h2("Position sizing")
    risk_inr = sleeve * risk_pct
    qty_by_risk = int(risk_inr // stop_dist)
    qty_by_cap = int(cap // entry)
    qty = max(0, min(qty_by_risk, qty_by_cap))
    bind = "notional cap" if qty == qty_by_cap and qty < qty_by_risk else (
        "dollar risk" if qty == qty_by_risk and qty < qty_by_cap else "tie")
    notional = qty * entry
    actual_risk = qty * R

    code_block(
        f"risk_inr     = ₹25,00,000 × {risk_pct_pretty}  = {fmt_inr(risk_inr)}\n"
        f"qty_by_risk  = floor( {fmt_inr(risk_inr)} / ₹30 )  = {qty_by_risk}\n"
        f"qty_by_cap   = floor( {fmt_inr(cap)} / ₹500 )       = {qty_by_cap}\n"
        f"qty          = min({qty_by_risk}, {qty_by_cap})  = {qty}   ← {bind} binds\n"
        f"notional     = {qty} × ₹500  = {fmt_inr(notional)}\n"
        f"actual_risk  = {qty} × ₹30   = {fmt_inr(actual_risk)}   ({actual_risk/sleeve*100:.2f}% of sleeve)"
    )
    if bind == "notional cap":
        para(
            f"Cap binds → actual risk {fmt_inr(actual_risk)} "
            f"({actual_risk/sleeve*100:.2f}%) is below the {risk_pct_pretty} target. "
            "Expected for low-ATR names — by design.",
            italic=True,
        )

    h2("Stop loss & R-multiple ladder")
    R1, R2 = partial_R
    t1_price = entry + R1 * R
    t2_price = entry + R2 * R
    partial_qty = max(1, int(qty * 33 / 100))
    residual_qty = qty - 2 * partial_qty
    code_block(
        f"initial_stop  = ₹500 − (2 × ₹15)  = ₹470.00\n"
        f"R             = ₹500 − ₹470       = ₹30 per share\n\n"
        f"1st partial:  +{R1}R  → ₹500 + ({R1} × ₹30) = ₹{t1_price:.2f}\n"
        f"2nd partial:  +{R2}R  → ₹500 + ({R2} × ₹30) = ₹{t2_price:.2f}\n\n"
        f"partial_qty   = floor({qty} × 33%) = {partial_qty} shares\n"
        f"residual_qty  = {qty} − 2 × {partial_qty} = {residual_qty} shares (trails)"
    )

    h2("Scenario A — Trend works (best case)")
    para(
        f"At +{R1}R (₹{t1_price:.2f}) → 1st partial; stop → breakeven. "
        f"At +{R2}R (₹{t2_price:.2f}) → 2nd partial; trail armed. "
        "Residual trails out at +9R (₹770)."
    )
    trail_exit = entry + 9 * R
    pnl_p1 = (t1_price - entry) * partial_qty
    pnl_p2 = (t2_price - entry) * partial_qty
    pnl_res = (trail_exit - entry) * residual_qty
    total_A = pnl_p1 + pnl_p2 + pnl_res
    table(
        ["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        [
            [f"1st partial at +{R1}R", f"₹{t1_price:.2f}", str(partial_qty), fmt_inr(pnl_p1), fmt_inr(pnl_p1)],
            [f"2nd partial at +{R2}R", f"₹{t2_price:.2f}", str(partial_qty), fmt_inr(pnl_p2), fmt_inr(pnl_p1 + pnl_p2)],
            ["Trail residual @ +9R", f"₹{trail_exit:.2f}", str(residual_qty), fmt_inr(pnl_res), fmt_inr(total_A)],
        ],
        [6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario A total: {fmt_inr(total_A)}  ({total_A/actual_risk:.2f}R)")

    h2("Scenario B — Pop then reverse")
    para(f"Hits +{R1}R, 1st partial fires, BE stop catches the pullback.")
    pnl_B_p1 = (t1_price - entry) * partial_qty
    total_B = pnl_B_p1
    table(
        ["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        [
            [f"1st partial at +{R1}R", f"₹{t1_price:.2f}", str(partial_qty), fmt_inr(pnl_B_p1), fmt_inr(pnl_B_p1)],
            ["Breakeven stop", "₹500.00", str(qty - partial_qty), fmt_inr(0), fmt_inr(total_B)],
        ],
        [6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario B total: {fmt_inr(total_B)}  ({total_B/actual_risk:.2f}R)")

    h2("Scenario C — Stopped cold (worst case)")
    pnl_C = -actual_risk
    table(
        ["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        [["Initial stop hit", "₹470.00", str(qty), fmt_inr(pnl_C), fmt_inr(pnl_C)]],
        [6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario C total: {fmt_inr(pnl_C)}  (−1.00R)")

    h2("Scenario D — Vertical pop (universal 25%)")
    para(
        "+25% within 15td → universal book 25% of initial qty; arms breakeven. "
        f"Trade then runs to +{R2}R and trails out at +9R."
    )
    uni_qty = max(1, int(qty * 25 / 100))
    pnl_uni = (625 - entry) * uni_qty
    remain = qty - uni_qty
    p1_q = min(partial_qty, remain); pnl_D_p1 = (t1_price - entry) * p1_q; remain -= p1_q
    p2_q = min(partial_qty, remain); pnl_D_p2 = (t2_price - entry) * p2_q; remain -= p2_q
    pnl_D_res = (trail_exit - entry) * remain
    total_D = pnl_uni + pnl_D_p1 + pnl_D_p2 + pnl_D_res
    table(
        ["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        [
            ["Universal 25% book", "₹625.00", str(uni_qty), fmt_inr(pnl_uni), fmt_inr(pnl_uni)],
            [f"1st partial at +{R1}R", f"₹{t1_price:.2f}", str(p1_q), fmt_inr(pnl_D_p1), fmt_inr(pnl_uni + pnl_D_p1)],
            [f"2nd partial at +{R2}R", f"₹{t2_price:.2f}", str(p2_q), fmt_inr(pnl_D_p2), fmt_inr(pnl_uni + pnl_D_p1 + pnl_D_p2)],
            ["Trail residual @ +9R", f"₹{trail_exit:.2f}", str(remain), fmt_inr(pnl_D_res), fmt_inr(total_D)],
        ],
        [6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario D total: {fmt_inr(total_D)}  ({total_D/actual_risk:.2f}R)")

    h2("Scenario E — Time stop (dead money)")
    para("25td later, price = ₹505 (inside [₹490, ₹510]). Time stop closes everything.")
    pnl_E = (505 - entry) * qty
    table(
        ["Milestone", "Price", "Qty", "P&L", "Cumulative P&L"],
        [["Time stop @ day 25", "₹505.00", str(qty), fmt_inr(pnl_E), fmt_inr(pnl_E)]],
        [6.0, 2.5, 2.0, 3.0, 3.5],
    )
    callout(f"Scenario E total: {fmt_inr(pnl_E)}  ({pnl_E/actual_risk:+.2f}R)")

    h2("Scenario summary for this tier")
    table(
        ["Scenario", "Total P&L", "R-multiple", "% of sleeve"],
        [
            ["A — Trend works", fmt_inr(total_A), f"+{total_A/actual_risk:.2f}R", f"{total_A/sleeve*100:+.2f}%"],
            ["B — Pop then reverse", fmt_inr(total_B), f"+{total_B/actual_risk:.2f}R", f"{total_B/sleeve*100:+.2f}%"],
            ["C — Stopped cold", fmt_inr(pnl_C), "−1.00R", f"{pnl_C/sleeve*100:+.2f}%"],
            ["D — Vertical pop", fmt_inr(total_D), f"+{total_D/actual_risk:.2f}R", f"{total_D/sleeve*100:+.2f}%"],
            ["E — Time stop", fmt_inr(pnl_E), f"{pnl_E/actual_risk:+.2f}R", f"{pnl_E/sleeve*100:+.2f}%"],
        ],
        [6.0, 3.5, 3.0, 3.0],
    )

    return {
        'tier': tier_name, 'qty': qty, 'notional': notional, 'risk': actual_risk,
        'A': total_A, 'B': total_B, 'C': pnl_C, 'D': total_D, 'E': pnl_E,
    }


# Now actually render all four tiers with proper section numbers.
# Discard the earlier render_tier call's section (it added section 3 already).
# We can't easily delete it, so we'll regenerate the doc from scratch.

# Recreate the doc cleanly:
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Algo Paper Trading — Tier-by-Tier Worked Examples")
r.bold = True; r.font.size = Pt(20)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Every number, every milestone, every tier — with actual rupees")
r.italic = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
doc.add_paragraph()

# Re-add sections 1 and 2
h1("1. The Mental Model")
para("Three numbers drive everything that follows. Memorize these:", bold=True)
code_block(
    "R          = entry_price − initial_stop      (per-share risk in ₹)\n"
    "stop_dist  = 2 × ATR_14                       (= 1R, by construction)\n"
    "SLEEVE     = ₹25,00,000                       (fixed paper capital)"
)
para(
    "Because stop_dist is fixed at 2×ATR, R is always equal to 2×ATR. So every "
    "R-multiple in the strategy translates directly into an ATR multiple of distance:"
)
table(
    ["R-multiple", "In ATR units", "Meaning"],
    [
        ["1R", "2 × ATR", "Per-share risk. Distance from entry to initial stop."],
        ["2R", "4 × ATR", "First partial target for T2/T3/T4."],
        ["3R", "6 × ATR", "First partial target for T1."],
        ["4R", "8 × ATR", "Second partial target for T2/T3/T4."],
        ["6R", "12 × ATR", "Second partial target for T1."],
    ],
    [2.5, 3.0, 11.5],
)

para(
    "Position size is constrained by two limits — whichever produces the SMALLER qty wins."
)
code_block(
    "qty_by_risk = floor( SLEEVE × risk_pct / stop_dist )    # dollar risk constraint\n"
    "qty_by_cap  = floor( cap / entry_price )                # notional cap constraint\n"
    "qty         = min(qty_by_risk, qty_by_cap)\n"
    "REJECT if qty × entry_price < ₹40,000  (position floor)"
)
table(
    ["Tier", "risk_pct", "Dollar risk @ ₹25L", "Notional cap"],
    [
        ["T1_MULTI_STRONG", "1.50%", "₹37,500", "₹3,00,000"],
        ["T2_STRONG_REG",   "1.00%", "₹25,000", "₹2,50,000"],
        ["T3_MULTI_REG",    "0.66%", "₹16,500", "₹1,50,000"],
        ["T4_RS_ACCEL",     "0.50%", "₹12,500", "₹1,00,000"],
    ],
    [4.5, 2.5, 4.0, 4.0],
)

h1("2. Shared Worked-Example Setup")
para(
    "Every tier uses the SAME underlying stock so differences are purely from risk_pct, "
    "cap, and the partial-R ladder."
)
code_block(
    "Hypothetical stock: XYZ.NS\n"
    "  D1 09:15 first tick     = ₹499.25\n"
    "  Slippage 15bps          → entry_price = ₹500.00\n"
    "  ATR_14 (from D0)        = ₹15.00       (3% of price — a typical mid/large cap)\n"
    "  stop_dist  (2 × ATR)    = ₹30.00\n"
    "  initial_stop            = ₹470.00\n"
    "  R                       = ₹30 per share"
)
para("Five exit scenarios per tier:", bold=True)
bullet("A — Trend works: hits both R-partials, residual trails up to +9R then stops.")
bullet("B — Pop then reverse: 1st partial fills, breakeven stop catches the pullback.")
bullet("C — Stopped cold: straight down to initial stop. Worst case.")
bullet("D — Vertical pop: +25% inside 15 trading days → universal 25% book triggers.")
bullet("E — Time stop: 25 td of dead money, returns inside [−2%, +2%].")

# Render tiers
t1_data = render_tier_numbered(3, "Tier T1", "Multi-Strong (highest conviction)", 0.015, "1.50%", 300_000, (3, 6))
t2_data = render_tier_numbered(4, "Tier T2", "Strong + Regular", 0.010, "1.00%", 250_000, (2, 4))
t3_data = render_tier_numbered(5, "Tier T3", "Multi-Regular", 0.0066, "0.66%", 150_000, (2, 4))
t4_data = render_tier_numbered(6, "Tier T4", "RS Acceleration (presignal-only)", 0.005, "0.50%", 100_000, (2, 4))

# ──────────────────────────────────────────────────────────────────────
# SIDE BY SIDE
# ──────────────────────────────────────────────────────────────────────
h1("7. Side-by-Side: Same Stock, Same Setup")
para(
    "Below is the exact same hypothetical XYZ trade (entry ₹500, ATR ₹15, R = ₹30) "
    "sized and run through all 5 scenarios for each tier. Read the columns left-to-right "
    "to see how tier choice changes the magnitude of outcomes."
)

table(
    ["Metric", "T1", "T2", "T3", "T4"],
    [
        ["Qty (shares)", t1_data['qty'], t2_data['qty'], t3_data['qty'], t4_data['qty']],
        ["Notional",   fmt_inr(t1_data['notional']), fmt_inr(t2_data['notional']),
                       fmt_inr(t3_data['notional']), fmt_inr(t4_data['notional'])],
        ["Actual ₹ risk", fmt_inr(t1_data['risk']), fmt_inr(t2_data['risk']),
                          fmt_inr(t3_data['risk']), fmt_inr(t4_data['risk'])],
        ["1st partial target", "₹590 (+3R)", "₹560 (+2R)", "₹560 (+2R)", "₹560 (+2R)"],
        ["2nd partial target", "₹680 (+6R)", "₹620 (+4R)", "₹620 (+4R)", "₹620 (+4R)"],
        ["A — Trend works",  fmt_inr(t1_data['A']), fmt_inr(t2_data['A']),
                              fmt_inr(t3_data['A']), fmt_inr(t4_data['A'])],
        ["B — Pop & reverse", fmt_inr(t1_data['B']), fmt_inr(t2_data['B']),
                               fmt_inr(t3_data['B']), fmt_inr(t4_data['B'])],
        ["C — Stopped cold",  fmt_inr(t1_data['C']), fmt_inr(t2_data['C']),
                               fmt_inr(t3_data['C']), fmt_inr(t4_data['C'])],
        ["D — Vertical pop",  fmt_inr(t1_data['D']), fmt_inr(t2_data['D']),
                               fmt_inr(t3_data['D']), fmt_inr(t4_data['D'])],
        ["E — Time stop",     fmt_inr(t1_data['E']), fmt_inr(t2_data['E']),
                               fmt_inr(t3_data['E']), fmt_inr(t4_data['E'])],
    ],
    [4.5, 3.0, 3.0, 3.0, 3.0],
)
para("Key observations:", bold=True)
bullet(
    "For this stock (low ATR, ₹500 price), the notional cap binds on ALL four tiers. "
    "That means actual risk per trade is well below the tier's risk_pct target — and "
    "differences between tiers in this example come almost entirely from the cap, "
    "not from the dollar-risk %.")
bullet(
    "T1 not only has the largest position but also a higher R-multiple ladder (+3R/+6R), "
    "so its winning scenarios are dramatically bigger than T4's. T1 Scenario A is roughly "
    "3× the rupee P&L of T4 Scenario A.")
bullet(
    "The worst case (C) scales linearly with position size. T4's worst loss is just over "
    "₹6,000 vs T1's ₹18,000 — exactly proportional to qty, which is the whole point of "
    "smaller positions for less-conviction tiers.")
bullet(
    "Scenario B (BE pullback) leaves you with a small win on the 1st partial and zero on "
    "the rest — a 'free look' outcome. Across many trades this dampens drawdown.")

# ──────────────────────────────────────────────────────────────────────
# HIGH-ATR SCENARIO
# ──────────────────────────────────────────────────────────────────────
h1("8. What if the stock is more volatile? (High-ATR case)")
para(
    "All numbers above assume ATR = 3% of price (mid/large cap behavior). For a higher-ATR "
    "name, stop_dist gets bigger, which means qty_by_risk drops AND the partial targets "
    "live further away in rupee terms. Let's run the same exercise on a more volatile "
    "stock:"
)
code_block(
    "Hypothetical stock: ZZZ.NS  (small/mid-cap, higher ATR)\n"
    "  entry_price             = ₹500.00\n"
    "  ATR_14                  = ₹40.00       (8% of price)\n"
    "  stop_dist  (2 × ATR)    = ₹80.00\n"
    "  initial_stop            = ₹420.00\n"
    "  R                       = ₹80 per share"
)
# High-ATR sizing
sleeve = 2_500_000; entry = 500.0; stop_dist_hi = 80.0; R_hi = 80.0
rows_hi = []
for tier, rp, cap in [("T1", 0.015, 300_000), ("T2", 0.010, 250_000), ("T3", 0.0066, 150_000), ("T4", 0.005, 100_000)]:
    qty_r = int((sleeve * rp) // stop_dist_hi)
    qty_c = int(cap // entry)
    qty = min(qty_r, qty_c)
    bind = "RISK" if qty_r <= qty_c else "CAP"
    notional = qty * entry
    actual_risk = qty * R_hi
    floor_ok = "OK" if notional >= 40_000 else "REJECTED (floor)"
    rows_hi.append([tier, qty_r, qty_c, qty, bind, fmt_inr(notional),
                    fmt_inr(actual_risk), floor_ok])

table(
    ["Tier", "qty_by_risk", "qty_by_cap", "Final qty", "Binds", "Notional", "Risk", "Position floor"],
    rows_hi,
    [1.8, 2.2, 2.2, 2.0, 1.6, 2.5, 2.0, 3.0],
)
para(
    "On a high-ATR name the DOLLAR RISK binds (qty_by_risk < qty_by_cap), and "
    "all four tiers end up at full risk %. The notional and qty are much smaller "
    "than in the ₹500/₹15-ATR case because the stop is so much wider. This is "
    "the intended behavior — volatility-adjusted sizing.",
    italic=True,
)
para(
    "Note: T4 on a high-ATR stock ends up with a very small position (~150 shares × ₹500 = "
    "₹78,000 notional). Still above the ₹40,000 floor, so it trades. But if entry_price "
    "were lower or risk_pct smaller, the floor could kill the trade entirely.",
    italic=True,
)

# ──────────────────────────────────────────────────────────────────────
# WHEN A TRADE IS REJECTED
# ──────────────────────────────────────────────────────────────────────
h1("9. When a Trade Gets Rejected at Sizing")
para(
    "Position floor = ₹40,000. The trade is rejected entirely if qty × entry_price < ₹40k. "
    "Common rejection paths:"
)
bullet(
    "High-priced low-ATR name where the cap forces a tiny qty. Example: ₹5,000 share, "
    "ATR ₹30. T4 cap = ₹1L → qty_by_cap = 20. Notional = 20 × ₹5,000 = ₹1,00,000 — "
    "actually fine. But for T4 you'd be at full cap on a single share count of 20.")
bullet(
    "Very high-ATR name where dollar risk forces qty into low single digits. Example: "
    "₹500 share, ATR ₹100 (20%). T4: risk_inr = ₹12,500. qty_by_risk = floor(12500 / 200) "
    "= 62. Notional = ₹31,000 < ₹40,000 → REJECT. Trade never gets placed.")
bullet(
    "Sector concentration. If the new trade's notional would push the sector total over "
    "25% × ₹25L = ₹6,25,000, the trade is rejected (regardless of sizing). This is "
    "checked AFTER sizing — the candidate is skipped and we move to the next.")

# ──────────────────────────────────────────────────────────────────────
# EXIT REASON CODES
# ──────────────────────────────────────────────────────────────────────
h1("10. Exit Reason Codes (what you'll see in the Closed Trades table)")
table(
    ["exit_reason", "Triggered by", "What happened"],
    [
        ["stop",            "Intraday (Railway) or EOD batch",
         "Initial stop hit BEFORE any partial or trail. Scenario C."],
        ["trail_stop",      "Intraday (Railway) or EOD batch",
         "Stop hit AFTER trail was armed (i.e. after 2nd partial). Scenario A's tail."],
        ["partials_full",   "EOD batch only",
         "Both partial targets filled in the same bar — no residual qty left."],
        ["time_stop",       "EOD batch",
         "25 trading days held, return inside [−2%, +2%]. Scenario E."],
        ["fill_expired",    "Railway fill job",
         "Pending row aged > 2 trading days with no live tick — never got filled."],
        ["fill_rejected_position_floor", "Railway fill job",
         "D1 open moved enough that sizing produced qty=0 or notional < ₹40k."],
        ["fill_rejected_missing_atr",    "Railway fill job",
         "entry_atr was null on the pending row (data quality issue — shouldn't happen)."],
    ],
    [4.5, 4.5, 8.0],
)

# Save
doc.save(OUT)
print(f"Wrote: {OUT}")
